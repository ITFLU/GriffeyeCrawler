#!/usr/bin/env python
#-*- coding:utf-8 -*-
"""
GRIFFEYE-ANALYZER
-----------------
Analysiert eine exportierte Dateiliste aus Griffeye pro Gerät & Kategorie
- Summiert die Bilder und Videos
- Fasst die Dateipfade zusammen und unterteilt diese in Cache- & Nicht-Cache-Pfade auf
- Ermittelt die Pfade mit den meisten Inhalten
- Ermittelt das prozentuelle Verhältnis im Browsercache und der übrigen Ablage
- Ermittelt die prozentuelle Verteilung der Dateierstellung im betroffenen Zeitraum
- Generiert eine Ergebnisdatei im TXT oder DOCX-Format

(c) 2021, Luzerner Polizei
Author:  Michael Wicki
Version: 09.04.2021
"""

import os
import sys
import json
import traceback
from datetime import datetime
# docx...
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL



class Device:
    def __init__(self, sourceid, category, initial_date):
        self.sourceid = sourceid
        self.categories = { category: Category(category, initial_date) }
        self.legal_count = 0
        self.illegal_count = 0

    def addDate(self, category, date):
        if category not in self.categories.keys():
            self.categories[category] = Category(category, date)
        else:
            self.categories[category].addDate(date)

    def addFile(self, category, path, mediatype, date):
        self.categories[category].addFile(path, mediatype, date)
        if category_legality.get(category, True):
            self.legal_count += 1
        else:
            self.illegal_count += 1

    def getSourceId(self):
        return self.sourceid

    def getCategory(self, category):
        if category not in self.categories.keys():
            return None
        else:
            return self.categories[category]

    def getCounts(self):
        """
        returns a tuple with total count, legal count & illegal count of the device
        """
        return (self.legal_count+self.illegal_count, self.legal_count, self.illegal_count)


class Category:
    def __init__(self, name, initial_date):
        self.name = name
        self.legality = category_legality.get(name, True)
        self.min_date = initial_date
        self.max_date = initial_date
        self.date_groups = {}
        self.pic_count = 0
        self.vid_count = 0
        self.tot_count = 0
        self.paths = {}
        self.cachepaths = {}

    def addDate(self, date):
        if date < self.min_date:
            self.min_date = date
        if date > self.max_date:
            self.max_date = date

    def addFile(self, path, mediatype, date):
        # increase counters
        self.tot_count += 1
        if mediatype == "Image":
            self.pic_count += 1
        if mediatype == "Video":
            self.vid_count += 1

        # increase path
        self.increasePath(path)

        # increase date
        self.increaseDate(date)

    def getCacheGroup(self, path):
        for k in known_cache_paths.keys():
            if k in path:
                return known_cache_paths[k]
        return None

    def increasePath(self, path):
        if path in self.cachepaths.keys():
            # path is in cachepath >> increase count
            self.cachepaths[path] += 1
        else:
            # path NOT in cachepath >> check for cache
            cache_group = self.getCacheGroup(path)
            if cache_group is not None:
                # path is cachepath
                if cache_group not in self.cachepaths.keys():
                    self.cachepaths[cache_group] = 1    # create
                else:
                    self.cachepaths[cache_group] += 1   # increase
            else:
                # path is NOT cachepath
                if path not in self.paths.keys():
                    self.paths[path] = 1    # create
                else:
                    self.paths[path] += 1   # increase

    def increaseDate(self, date):
        year = int(date[6:10])
        if year not in self.date_groups.keys():
            self.date_groups[year] = 1  # create
        else:
            self.date_groups[year] += 1 # increase

    def getDateRange(self):
        return self.min_date.strftime("%d.%m.%Y")+" - "+self.max_date.strftime("%d.%m.%Y")

    def getDateRangeDays(self):
        return (self.max_date - self.min_date).days + 1

    def getCounts(self):
        """
        returns a tuple with total count, picture count & video count of the category
        """
        return (self.tot_count, self.pic_count, self.vid_count)
    
    def getCountsString(self):
        """
        returns a string with formatted picture- & videos-count
        """
        result = ""
        # pictures
        if self.pic_count > 0:
            result += "{} ".format(self.pic_count)
            if self.pic_count > 1:
                result += "Bilder"
            else:
                result += "Bild"
            if self.vid_count > 0:
                result += ", "
        # videos
        if self.vid_count > 0:
            result += "{} ".format(self.vid_count)
            if self.vid_count > 1:
                result += "Videos"
            else:
                result += "Video"
        return result

    def getGroupedDates(self):
        """
        returns a string with the percentage of illegals files per year
        """
        result = ""
        for year in sorted(self.date_groups.keys()):
            # calculate percentage of total files
            perc = (self.date_groups[year]/self.tot_count)*100
            result = result+"{}: {:.0f}%, ".format(year, perc)
        return result[:-2] # kill last ', '

    def getBrowserCacheSum(self):
        sum = 0
        for c in self.cachepaths:
            if c in browser_names:
                sum += self.cachepaths[c]
        return sum


class ColumnNotFoundException(Exception):
    """
    Error bei nicht gefundenem Spaltennamen
    """
    def __init__(self, columnname):
        self.message = "Spalte '{}' wurde nicht gefunden".format(columnname)



def progress(count, total, status=''):
    """
    handling of the progressbar
    """
    bar_len = 60
    filled_len = int(round(bar_len * count / float(total)))
    percents = round(100.0 * count / float(total), 1)
    bar = '#' * filled_len + '-' * (bar_len - filled_len)

    sys.stdout.write('[%s] %s%s ...%s\r' % (bar, percents, '%', status))
    sys.stdout.flush()


def getLinecount(filename):
    """
    count total lines for progressbars
    """
    counter = -1
    file_input = open(filename, "r", encoding="utf16")
    for line in file_input:
        counter += 1
        if counter == 0:
            continue
    file_input.close()
    return counter

def getTitleString(title, symbol="-", length=70):
    """
    creates a titleline with centered text
    """
    half_length = (length//2)-1  # including blank
    half_title = (len(title)//2)-1  # including blank
    symbol_count = half_length-half_title
    addition = ""
    if half_title%2 > 0 or half_length%2 > 0:
        addition = symbol
    return symbol*symbol_count+" "+title+" "+symbol*symbol_count+addition

def shortenPath(path):
    """
    shortens the filepath by the first two directories
    """
    first = path[path.find(os.path.sep)+1:]
    return first[first.find(os.path.sep)+1:]

def checkColumns(header):
    """
    check for needed columns & fill columnindex-dictionary for column access with columnname
    """
    cols = header.split(';')
    for c in config["needed_columns"]:
        if c["columnname"] in cols:
            column_index[c["key"]] = cols.index(c["columnname"])
        else:
            raise ColumnNotFoundException(c["columnname"])

def analyzeFile(filename):
    """
    - check for needed columns & fill columnindex-dictionary
    - collect devices & fill devicelist
    - detect min & max date for daterange (per device)
    """
    counter = -1
    file_input = open(filename, "r", encoding="utf16")
    for line in file_input:
        counter += 1
        if counter == 0:
            # csv-header...
            checkColumns(line)
            continue

        # csv-entry...
        # get device & date from csv
        data = line.split(";")
        data_category = data[column_index['col_category']]
        data_date = data[column_index['col_date']]
        current_date = datetime.strptime(data_date[0:10], "%d.%m.%Y")
        data_device = data[column_index['col_device']]
        # check for device or create it when needed
        if data_device not in devices.keys():
            devices[data_device] = Device(data_device, data_category, current_date)
        else:
            devices[data_device].addDate(data_category, current_date)
        # update progressbar
        progress(counter, linecount)

    file_input.close()
    return counter

def writeOutputfileTxt():
    file_result = open(result_filename,"w", encoding="utf-8")
    # write results of file-analysis
    file_result.write("GRIFFEYE-ANALYZER - Ergebnis vom {}\n".format(datetime.now().strftime("%d.%m.%Y")))
    file_result.write("="*43+"\n")
    file_result.write("Analysierte Datei:     {}\n".format(input_filename))
    file_result.write("Anzahl Datensätze:     {}\n".format(linecount))
    file_result.write("\n")

    # write results of devices
    counter = 0
    for d in devices:
        counter += 1
        file_result.write("\n{}\n".format(getTitleString(d, "=")))
        file_result.write("{} Dateien (Legal: {}, Illegal: {})".format(devices[d].getCounts()[0], devices[d].getCounts()[1], devices[d].getCounts()[2]))
        file_result.write("  >>  {:.2f}% illegal\n".format((devices[d].getCounts()[2]/devices[d].getCounts()[0])*100))
        for c in sorted(category_sort.keys()):
            if category_sort[c] not in devices[d].categories:
                continue

            cat = devices[d].getCategory(category_sort[c])
            file_result.write("\n{}\n".format(getTitleString(cat.name, "\u0387")))
            # count & mediatype
            file_result.write("Menge/Dateityp:\t")
            file_result.write("{}\n".format(cat.getCountsString()))
            if category_sort[c] != "Legale Pornographie":
                # daterange
                file_result.write("Erstellung auf Datenträger:\t{}".format(cat.getDateRange()))
                file_result.write("\n")
                # timeline
                file_result.write("Verteilung im Zeitraum:\t{}".format(cat.getGroupedDates()))
                file_result.write("\n")
                # proportion storage <-> browser cache
                perc = (cat.getBrowserCacheSum()/cat.getCounts()[0])*100
                file_result.write("Anteil Browsercache:\t{:.0f}%".format(perc))
                file_result.write("\n")
                # paths
                file_result.write("Speicherorte:\t")
                file_result.write("\n")
                # show top-paths
                i = 0
                for k in sorted(cat.paths, key=cat.paths.get, reverse=True):
                    i += 1
                    if i > config["result"]["number_of_showed_paths"]:
                        break;
                    file_result.write("- {}\n".format(shortenPath(k)))
                # if available, write other caches
                if len(cat.cachepaths)>0:
                    file_result.write("    > Caches <\n")
                    for k in cat.cachepaths:
                        file_result.write("- {}\n".format(k))

        file_result.write("\n")

        # update progressbar
        progress(counter, len(devices))

    file_result.close()

def writeOutputfileDocx():
    text_fontname = "Arial"
    text_fontsize = Pt(11)
    table_fontsize = Pt(8)
    table_rowheight = Pt(14)

    document = Document()
    # write results of file-analysis
    document.add_heading("GRIFFEYE-ANALYZER - Ergebnis vom {}".format(datetime.now().strftime("%d.%m.%Y")), 1)
    p = document.add_paragraph()
    run = p.add_run("Analysierte Datei:\t{}\nAnzahl Datensätze:\t{}".format(input_filename, linecount))
    run.font.name = text_fontname
    run.font.size = text_fontsize

    # write results of devices
    counter = 0
    for d in devices:
        counter += 1
        document.add_heading(d, 2)
        p = document.add_paragraph()
        run = p.add_run("{} Dateien (Legal: {}, Illegal: {})  >>  {:.2f}% illegal".format(devices[d].getCounts()[0], devices[d].getCounts()[1], devices[d].getCounts()[2], (devices[d].getCounts()[2]/devices[d].getCounts()[0])*100))
        run.font.name = text_fontname
        run.font.size = text_fontsize
        run.italic = True

        for c in sorted(category_sort.keys()):
            if category_sort[c] not in devices[d].categories:
                continue

            cat = devices[d].getCategory(category_sort[c])
            # write table...
            table = document.add_table(rows=1, cols=2, style="Table Grid")
            # format header
            hdr_cells = table.rows[0].cells
            # cell merging
            hdr_cells[0].merge(hdr_cells[1])
            hdr_cells[0].text = cat.name
            # background color
            cellprop = hdr_cells[0]._tc.get_or_add_tcPr()
            cellshade = OxmlElement("w:shd")
            cellshade.set(qn("w:fill"), "#CCCCCC")
            cellprop.append(cellshade)
            # row alignment
            table.rows[0].height = table_rowheight
            hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # font
            run = hdr_cells[0].paragraphs[0].runs[0]
            run.font.name = text_fontname
            run.font.size = table_fontsize
            run.font.bold = True

            # fill data...
            # count & mediatype
            row_cells = table.add_row().cells
            row_cells[0].text = "Menge/Dateityp:"
            row_cells[1].text = cat.getCountsString()
            if category_sort[c] != "Legale Pornographie":
                # daterange
                row_cells = table.add_row().cells
                row_cells[0].text = "Erstellung auf Datenträger:"
                row_cells[1].text = "{}".format(cat.getDateRange())
                # timeline
                row_cells = table.add_row().cells
                row_cells[0].text = "Verteilung im Zeitraum:"
                row_cells[1].text = "{}".format(cat.getGroupedDates())
                # proportion storage <-> browser cache
                row_cells = table.add_row().cells
                row_cells[0].text = "Anteil Browsercache:"
                perc = (cat.getBrowserCacheSum()/cat.getCounts()[0])*100
                row_cells[1].text = "{:.0f}%".format(perc)
                # paths
                row_cells = table.add_row().cells
                row_cells[0].text = "Speicherort(e):"
                # show top-paths
                rows = ""
                i = 0
                for k in sorted(cat.paths, key=cat.paths.get, reverse=True):
                    i += 1
                    if i > config["result"]["number_of_showed_paths"]:
                        break
                    if i > 1:
                        rows += "\n"
                    rows += "- {}".format(shortenPath(k))
                # # if available, write other caches
                # if len(cat.cachepaths)>0:
                #     file_result.write("    > Caches <\n")
                #     for k in cat.cachepaths:
                #         file_result.write("- {}\n".format(k))
                row_cells[1].text = rows
            
            # format table
            for row in table.rows[1:]:
                i=-1
                row.height = table_rowheight
                for cell in row.cells:
                    i+=1
                    # cell alignment
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # font
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = text_fontname
                    run.font.size = table_fontsize
                    if i == 0:
                        # name-column bold
                        run.font.bold = True
            document.add_paragraph().paragraph_format.space_after = Pt(0)

        progress(counter, len(devices))
        document.save(result_filename)
    
def writePathDetails():
    """
    creates the outputfile (txt) with detailed information
    """
    file_result = open(config["result"]["pathdetails_name"],"w", encoding="utf-8")
    # write results of file-analyze
    file_result.write("GRIFFEYE-ANALYZER - Pfad-Details vom {}\n".format(datetime.now().strftime("%d.%m.%Y")))
    file_result.write("="*47+"\n")
    file_result.write("Analysierte Datei:     {}\n".format(input_filename))
    file_result.write("Anzahl Datensätze:     {}\n".format(linecount))
    file_result.write("\n")

    # write results of devices
    counter = 0
    for d in devices:
        counter += 1
        file_result.write("\n{}\n".format(getTitleString(d, "=")))
        file_result.write("{} Dateien (Legal: {}, Illegal: {})".format(devices[d].getCounts()[0], devices[d].getCounts()[1], devices[d].getCounts()[2]))
        file_result.write("  >>  {:.2f}% illegal\n".format((devices[d].getCounts()[2]/devices[d].getCounts()[0])*100))
        for c in sorted(category_sort.keys()):
            if category_sort[c] not in devices[d].categories:
                continue

            cat = devices[d].getCategory(category_sort[c])
            file_result.write("\n{}\n".format(getTitleString(cat.name, "\u0387")))
            # count & mediatype
            file_result.write("Menge/Dateityp:\t")
            file_result.write("{}\n".format(cat.getCountsString()))

            # daterange
            file_result.write("Erstellung auf Datenträger:\t{}".format(cat.getDateRange()))
            file_result.write("\n")
            # timeline
            file_result.write("Verteilung im Zeitraum:\t{}".format(cat.getGroupedDates()))
            file_result.write("\n")
            # proportion storage <-> browser cache
            perc = (cat.getBrowserCacheSum()/cat.getCounts()[0])*100
            file_result.write("Anteil Browsercache:\t{:.0f}% >>> (Total: {}, Browsercache: {})".format(perc, cat.getCounts()[0], cat.getBrowserCacheSum()))
            file_result.write("\n")
            # paths
            file_result.write("Speicherorte:\t")
            file_result.write("\n")

            # show paths
            for k in sorted(cat.paths, key=cat.paths.get, reverse=True):
                file_result.write("- {} >>> {}\n".format(k, str(cat.paths[k])))
            # if available, write other caches
            if len(cat.cachepaths)>0:
                file_result.write("    > Caches <\n")
                for k in cat.cachepaths:
                    file_result.write("- {} >>> {}\n".format(k, str(cat.cachepaths[k])))

        file_result.write("\n")
        # update progressbar
        progress(counter, len(devices))

    file_result.close()


# init
column_index = {}
devices = {}
linecount = 0

try:
    print("===== GRIFFEYE-ANALYZER =====")

    # read configurations
    with open('config.json') as c:
        data = c.read()
    config = json.loads(data)
    input_filename = config["input"]["filename"]
    result_filename = config["result"]["filename"]
    category_legality = {}
    category_sort = {}
    for cat in config["categories"]:
        category_legality[cat["name"]] = cat["legality"]
        category_sort[cat["sort"]] = cat["name"]
    known_cache_paths = {}
    browser_names = []
    for cac in config["caches"]:
        known_cache_paths[cac["path"]] = cac["name"]
        if cac["is_browser"] and cac["name"] not in browser_names:
            browser_names.append(cac["name"])

    # ask for names & options
    input_filename = input("Name des Input-CSV (Default: {}) > ".format(input_filename)) or input_filename
    result_filename = input("Name der Ergebnisdatei (Default: {}) [.txt, .docx] > ".format(result_filename)) or result_filename
    result_format = "txt"
    if ".docx" in result_filename:
        result_format = "docx"
    if ".txt" in result_filename:
        result_format = "txt"
    print()

    # get linecount for progressbar
    linecount = getLinecount(input_filename)

    # analyze file
    print("Analysiere Datei '{}'...".format(input_filename))
    analyzeFile(input_filename)
    print()

    # process data
    print("Verarbeite Datensätze...")
    file_input = open(input_filename, "r", encoding="utf16")
    result = ""
    counter = 0
    for line in file_input:
        counter += 1
        # ignore csv-header
        if counter == 1:
            continue

        # get data from file
        column = line.split(";")
        data_category = column[column_index['col_category']]
        data_path = column[column_index['col_path']]
        data_type = column[column_index['col_type']]
        data_date = column[column_index['col_date']]
        data_device = column[column_index['col_device']]
        # add data to device
        device = devices[data_device]
        device.addFile(data_category, data_path, data_type, data_date)
        # update progressbar
        progress(counter, linecount)

    file_input.close()
    print()

    # write output-files
    print("Schreibe Ergebnisdatei...")
    if result_format == "txt":
        writeOutputfileTxt()
    elif result_format == "docx":
        writeOutputfileDocx()
    else:
        print("bummmm...")
    if config["result"]["generate_pathdetails"]:
        writePathDetails()

    print()
    print()
    print("ERLEDIGT! {} Datensätze verarbeitet (siehe '{}')".format(counter, result_filename))

except ColumnNotFoundException as exp:
    print("ERROR: Verarbeitung abgebrochen!")
    print(">", exp.message)
except FileNotFoundError as exp:
    print("ERROR: Verarbeitung abgebrochen!")
    print("> Datei '{}' nicht gefunden".format(exp.filename))
except KeyError as exp:
    print("ERROR: Verarbeitung abgebrochen!")
    print("> Konfiguration {} nicht gefunden".format(exp))
except Exception as exp:
    print("ERROR: Verarbeitung abgebrochen!")
    traceback.print_exc()
    # print(">", type(exp))
    # print(exp)

print()
input()

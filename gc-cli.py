#!/usr/bin/env python
#-*- coding:utf-8 -*-
"""
GRIFFEYE-CRAWLER CLI
--------------------
Analyzes an exported file list of Griffeye per device & category

(c) 2023, Luzerner Polizei
Author:  Michael Wicki
"""
version = "1.3"

import argparse

import os
import sys
import json
import traceback
from datetime import datetime
# docx...
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL


MEDIATYPE_IMAGE = "Image"
MEDIATYPE_VIDEO = "Video"
MEDIATYPE_IGNORE = "ignore"


class Device:
    """
    class for the data per device
    """
    def __init__(self, sourceid):
        self.sourceid = sourceid
        self.categories = {}
        self.legal_count = 0
        self.illegal_count = 0

    def add_file(self, category, path, mediatype, date, hash):
        if category not in self.categories.keys():
            self.categories[category] = Category(category)
        self.categories[category].add_file(path, mediatype, date, hash)
        
        # increase legal/illegal count
        if category_legality.get(category, True):
            self.legal_count += 1
        else:
            self.illegal_count += 1

    def add_separate_thumb(self, category, path, mediatype, hash):
        if category not in self.categories.keys():
            self.categories[category] = Category(category)
        self.categories[category].add_separate_thumb(path, mediatype, hash)

    def get_sourceid(self):
        return self.sourceid

    def get_categories(self):
        return self.categories

    def get_category(self, category):
        if category not in self.categories.keys():
            return None
        else:
            return self.categories[category]

    def get_counts(self):
        """
        returns a tuple with total count, legal count & illegal count of the device
        """
        return (self.legal_count+self.illegal_count, self.legal_count, self.illegal_count)


class Category:
    """
    class for the data per category per device (included in Device)
    """
    def __init__(self, name):
        self.name = name
        self.legality = category_legality.get(name, True)
        self.visible = category_visibilty.get(name, True)
        self.min_date = empty_date
        self.max_date = empty_date
        self.year_groups = {}
        self.pic_count = 0
        self.vid_count = 0
        self.tot_count = 0
        self.paths = {} # paths which are not in a cache (path: Path)
        self.caches = {} # caches (name: Cache)
        self.separate_thumbs = {} # thumbcaches if separated > --includethumbs integrates it in self.paths (path: Path)
        self.separate_thumbs_hashes = set()
        self.pic_hashes = set()
        self.vid_hashes = set()

    def add_file(self, path, mediatype, date, hash):
        # increase counters & add hash to 'hashes' (>> deduplicates itself)
        self.tot_count += 1
        if mediatype == MEDIATYPE_IMAGE:
            self.pic_count += 1
            self.pic_hashes.add(hash)
        if mediatype == MEDIATYPE_VIDEO:
            self.vid_count += 1
            self.vid_hashes.add(hash)

        self.recalculate_daterange(date)
        self.increase_path(path, mediatype)
        self.increase_date(date)

    def add_separate_thumb(self, path, mediatype, hash):
        if path not in self.separate_thumbs.keys():
            self.separate_thumbs[path] = Path(path, mediatype)
        else:
            self.separate_thumbs[path].increase_count(mediatype)
        self.separate_thumbs_hashes.add(hash)

    def get_separate_thumbs_total(self):
        sum = 0
        for path in self.separate_thumbs.values():
            sum += path.count_total
        return sum
    
    def get_separate_thumbs_total_unique(self):
        return len(self.separate_thumbs_hashes)

    def merge(self, merge_cat):
        """ merges another category to this one for the device totals """
        # merge daterange
        self.recalculate_daterange(merge_cat.min_date)
        self.recalculate_daterange(merge_cat.max_date)
        # merge counts
        self.tot_count += merge_cat.get_counts()[0]
        self.pic_count += merge_cat.get_counts()[1]
        self.vid_count += merge_cat.get_counts()[2]
        # merge dategroups
        for item in merge_cat.year_groups.items():
            key = item[0]
            value = item[1]
            if key not in self.year_groups.keys():
                self.year_groups[key] = value  # create
            else:
                self.year_groups[key] += value # increase
        # merge paths
        for path in merge_cat.paths.keys():
            if path not in self.paths.keys():
                self.paths[path] = merge_cat.paths[path]    # create
            else:
                self.paths[path].increase_object(merge_cat.paths[path])   # increase
        # merge separate_thumbs
        for path in merge_cat.separate_thumbs.keys():
            if path not in self.separate_thumbs.keys():
                self.separate_thumbs[path] = merge_cat.separate_thumbs[path]    # create
            else:
                self.separate_thumbs[path].increase_object(merge_cat.separate_thumbs[path])   # increase
        # merge caches
        for merge_cache in merge_cat.caches.values():
            for pathname in merge_cache.paths.keys():
                cache = self.get_cache(pathname)
                if cache is not None:
                    cache.add_path_object(merge_cache.paths[pathname])
        # merge hashes
        self.pic_hashes.update(merge_cat.pic_hashes)
        self.vid_hashes.update(merge_cat.vid_hashes)
        self.separate_thumbs_hashes.update(merge_cat.separate_thumbs_hashes)

    def recalculate_daterange(self, date):
        if date != empty_date and date != unix_date:
            if self.min_date == empty_date or date < self.min_date:
                self.min_date = date
            if self.max_date == empty_date or date > self.max_date:
                self.max_date = date

    def increase_path(self, path, mediatype):
        cache = self.get_cache(path)
        if cache is not None:
            cache.add_path(path, mediatype)
        else:
            if path not in self.paths.keys():
                self.paths[path] = Path(path, mediatype)    # create
            else:
                self.paths[path].increase_count(mediatype)   # increase

    def increase_date(self, date):
        year = int(date.strftime("%Y"))
        if year == 1 or year == 1970: # no date or unix date
            year = 9999
        if year not in self.year_groups.keys():
            self.year_groups[year] = 1  # create
        else:
            self.year_groups[year] += 1 # increase

    def get_cache(self, path):
        for cache in self.caches.values():
            if path in cache.paths.keys():
                # path exists already in cache
                return cache

        # path not found in cache > check for matching patterns
        for group in known_cache_names.values():
            for pattern in group.patterns:
                if pattern in path:
                    # path matches a cache pattern
                    if group.name in self.caches.keys():
                        return self.caches[group.name]
                    # cache exists not yet
                    cache = self.Cache(group)
                    self.caches[group.name] = cache
                    return cache
        return None
    
    def get_date_range(self):
        min = self.min_date.strftime(date_format)
        max = self.max_date.strftime(date_format)
        if min == empty_date:
            min = labels['undefined']
        if max == empty_date:
            max = labels['undefined']
        return (min, max)
    
    def get_date_range_string(self):
        if self.min_date == empty_date or self.max_date == empty_date:
            return labels['undefined']
        return self.min_date.strftime(date_format)+" - "+self.max_date.strftime(date_format)

    def get_unique_counts(self):
        """ returns a tuple with total count, picture count & video count of binary unique files (based on the hash) """
        return (len(self.pic_hashes.union(self.vid_hashes)), len(self.pic_hashes), len(self.vid_hashes))

    def get_counts(self):
        """ returns a tuple with total count, picture count & video count of the category """
        return (self.tot_count, self.pic_count, self.vid_count)
    
    def get_counts_string(self):
        """ returns a string with formatted picture- & videos-count """
        result = ""
        # pictures
        if self.pic_count > 0:
            result += f"{self.pic_count} "
            if self.pic_count > 1:
                result += labels['pictures']
            else:
                result += labels['picture']
            # binary unique
            result += f" ({len(self.pic_hashes)})"
            if self.vid_count > 0:
                result += ", "
        # videos
        if self.vid_count > 0:
            result += f"{self.vid_count} "
            if self.vid_count > 1:
                result += labels['videos']
            else:
                result += labels['video']
            # binary unique
            result += f" ({len(self.vid_hashes)})"
        if result == "":
            return "0"
        return result

    def get_grouped_years(self):
        """ returns a string with the percentage of illegal files per year """
        result = ""
        for year in sorted(self.year_groups.keys()):
            # calculate percentage of total files
            perc = (self.year_groups[year]/self.tot_count)*100
            if year == 9999:
                year = labels['undefined_short']
            perc_str = "{:.0f}%".format(perc)
            if round(perc, 0) == 0 and perc > 0:
                perc_str = "<1%"
            result = result+"{}: {}, ".format(year, perc_str)
        if result == "":
            return "-"
        return result[:-2] # kill last ', '

    def get_browsercache_total(self):
        """ returns the total count of browsercache files """
        sum = 0
        for c in self.caches.values():
            if c.group.is_browser:
                sum += c.count
        return sum

    def get_browsercache_sums(self):
        """ returns a dict with counts as Path object (value) for the specific browsers (key) """
        result = {}
        for c in self.caches.values():
            if c.group.is_browser:
                result[c.name] = Path(c.name, MEDIATYPE_IGNORE)
        return result

    def get_thumbcache_sum(self):
        """ returns the count of thumbcache files """
        sum = 0
        for c in self.caches.values():
            if c.group.is_thumbcache:
                sum += c.count
        return sum

    def get_thumbcache_obj(self):
        """ returns the thumbcache count as Path object """
        return Path(self.get_thumbcache_sum(), MEDIATYPE_IGNORE)
    
    class Cache:
        """
        inner class of Category for the data of containing cache paths (based on CacheGroup)
        """
        def __init__(self, group):
            self.name = group.name
            self.group = group
            self.paths = {} # path: Path
            self.count = 0
        
        def add_path(self, path, mediatype=MEDIATYPE_IGNORE):
            if path not in self.paths:
                self.paths[path] = Path(path, mediatype)
            else:
                self.paths[path].increase_count(mediatype)
            self.count += 1

        def add_path_object(self, path_obj):
            if path_obj.path not in self.paths:
                self.paths[path_obj.path] = path_obj
            else:
                self.paths[path_obj.path].increase_object(path_obj)
            self.count += path_obj.count_total


class Path:
    """
    class for the counts of files (total, picture, video) in a specific path
    """
    def __init__(self, path, mediatype):
        self.path = path
        self.count_total = 0
        self.count_pic = 0
        self.count_vid = 0
        self.show_details = False if mediatype==MEDIATYPE_IGNORE else True
        self.increase_count(mediatype)
    
    def increase_count(self, mediatype):
        self.count_total += 1
        if mediatype == MEDIATYPE_IMAGE:
            self.count_pic += 1
        if mediatype == MEDIATYPE_VIDEO:
            self.count_vid += 1
        
    def increase_object(self, path_obj):
        self.count_total += path_obj.count_total
        self.count_pic += path_obj.count_pic
        self.count_vid += path_obj.count_vid


class CacheGroup:
    """
    class for all existing cashes based on config.json (basic for Category/Case)
    """
    def __init__(self, name, is_browser, is_thumbcache):
        self.name = name
        self.is_browser = is_browser
        self.is_thumbcache = is_thumbcache
        self.patterns = []

    def add_pattern(self, pattern):
        if not pattern in self.patterns:
            self.patterns.append(pattern)


class PathNotFoundException(Exception):
    """ error in case of a path not found """
    def __init__(self, path):
        self.message = f"Path '{path}' not found"

class ColumnNotFoundException(Exception):
    """ error in case of a column name not found """
    def __init__(self, columnname):
        self.message = f"Column '{columnname}' not found"

class SeparatorNotFoundException(Exception):
    """ error in case of the column separator could not be detected """
    def __init__(self):
        self.message = f"Column separator could not be found... Please use option -s"

class LanguageNotFoundException(Exception):
    """ error in case of the language labels could not be detected """
    def __init__(self, language):
        self.message = f"Language '{language}' could not be found..."

class LineNotValidException(Exception):
    """ error in case of a csv-entry with ; in a field without " around it """
    def __init__(self, linenumber):
        self.message = f"Line '{linenumber}' is not valid"



def configure_argparse():
    global args
    parser = argparse.ArgumentParser(prog="gc-cli", 
                                     description="Commandline version of 'GriffeyeCrawler'\nAnalyze an exported filelist of Griffeye", 
                                     formatter_class=argparse.RawTextHelpFormatter,
                                     epilog='''\
Example of use
- JSON with dates from datefields prioritized as follows: 'exif dates' then 'last write time' then 'created date'
    python gc-cli.py --date "exif: createdate,last write time,created date" -f json metadata.csv
- DOCX in english excluding files in pathes including the texts 'unallocated' and 'unwantedfolder' in the pathname
    python gc-cli.py --exclude unallocated,unwantedfolder -l en_us metadata.csv
- JSON with new name in subfolder without details file but with max. 10 most common paths
    python gc-cli.py -o mysubfolder/mynew.json -n 10 --nodetails metadata.csv''')
    parser.version=version
    parser.add_argument("file", type=str, help="export csv of Griffeye")    
    parser.add_argument("-v", "--version", action="version")
    parser.add_argument("-o", metavar="output", action="store", type=str, 
                        help='''\
defines the output path/filename
could be only a path or can include a filename too
(default: input directory and input filename with the extension of the format)
defines the format too based on the file extension and overwrites -f''')
    parser.add_argument("-f", metavar="format", action="store", type=str, 
                        help=f'''\
defines the output format
overwritten by -o if a file extension is defined
possible values: {", ".join(map(str,valid_formats))} (default: {default_format})''')
    parser.add_argument("-l", metavar="language", action="store", type=str, 
                        help='''\
language for output documents (only partially for json) in locale format (e.g. en_US, de_DE)
if locale is not found, only the first part of the locale is checked (e.g. en, de)
languages are based on labels.json
(default from config.json, result/number_of_showed_paths)''')
    parser.add_argument("-n", metavar="number", action="store", type=int, help="number of paths to show per category")
    parser.add_argument("-s", metavar="separator", action="store", type=str, 
                        help='''\
defines the column separator
(default: automatically detected > comma or semicolon by Griffeye)''')
    parser.add_argument("-d", metavar="dateformat", action="store", type=str, help=f'''\
defines the format of the input date with format codes > see python help for more details
%%d  Day of the month (e.g. 01)
%%m  Month (e.g. 12)
%%y  Year without century (e.g. 23)
%%Y  Year with century (e.g. 2023)
needs to be wrapped in quotes if it contains a space
(default from config.json, input/date_format)''')
    parser.add_argument("--date", metavar="dates", action="store", type=str, 
                        help='''\
list of datefields to get the dates from separated by comma without space (case insensitive)
if a date is empty (01.01.0001, 01.01.1970 or '') the next field in the list is checked
needs to be wrapped in quotes if it contains a space
(default from config.json, needed_columns/col_date & other/alternative_date_column)''')
    parser.add_argument("--exclude", metavar="path", action="store", type=str, 
                        help='''\
list of textparts in the filepath field to be excluded from the analysis
separated by comma without space (case insensitive)
needs to be wrapped in quotes if it contains a space''')
    parser.add_argument("--nodetails", action="store_true", help="don't generate the pathdetails file")
    parser.add_argument("--includethumbs", action="store_true", help="include thumbcaches in the process (counts & dateranges) instead of listing them separately")
    args = parser.parse_args()

def progress(count, total, status=''):
    """ handling of the progressbar """
    bar_len = 60
    filled_len = int(round(bar_len * count / float(total)))
    percents = round(100.0 * count / float(total), 1)
    bar = '#' * filled_len + '-' * (bar_len - filled_len)

    sys.stdout.write('[%s] %s%s ...%s\r' % (bar, percents, '%', status))
    sys.stdout.flush()

def get_linecount(filename):
    """ count total lines for progressbars """
    counter = -1
    file_input = open(filename, "r", encoding=input_encoding)
    for line in file_input:
        counter += 1
        if counter == 0:
            continue
    file_input.close()
    return counter

def get_titlestring(title, symbol="-", length=70):
    """ creates a titleline with centered text """
    half_length = (length//2)-1  # including blank
    half_title = (len(title)//2)-1  # including blank
    symbol_count = half_length-half_title
    addition = ""
    if half_title%2 > 0 or half_length%2 > 0:
        addition = symbol
    return symbol*symbol_count+" "+title+" "+symbol*symbol_count+addition

def get_browser_percent(browser_count, total_count):
    if total_count==0:
        return "0%"
    perc = (browser_count/total_count)*100
    if round(perc, 0) == 0 and perc > 0:
        return "<1%"
    return "{:.0f}%".format(perc)

def shorten_path(path):
    """ shortens the filepath by the first two directories """
    first = path[path.find(os.path.sep)+1:]
    return first[first.find(os.path.sep)+1:]

def get_date_field(data):
    has_unix_date = False
    for i in column_index.keys():
        # ignore non-date-fields
        if not i.startswith("col_date"):
            continue
        # ignore empty fields ''
        if len(data[column_index[i]].strip()) == 0:
            continue

        date_obj = datetime.strptime(data[column_index[i]][0:10], date_format)
        # ignore empty dates '01.01.0001' > try next date (datefields_list is integrated...)
        if date_obj == empty_date:
            continue
        # ignore unix dates '01.01.1970'
        if date_obj == unix_date:
            has_unix_date = True
            continue
        return date_obj
    
    return unix_date if has_unix_date else empty_date
    # if has_unix_date:
    #     return unix_date
    # return empty_date

def is_thumbcache(path):
    for group in known_cache_names.values():
        if not group.is_thumbcache:
            continue
        for pattern in group.patterns:
            if pattern in path:
                return True
    return False

def detect_separator(header):
    """ detect the csv separator (, or ;) """
    global csv_separator
    if header.find(',') > -1:
        csv_separator = ","
    elif header.find(';') > -1:
        csv_separator = ";"
    else:
        raise SeparatorNotFoundException()

def check_columns(header):
    """ check for needed columns & fill columnindex-dictionary for column access with columnname """
    cols = header.strip('\n').split(csv_separator)
    for c in config["needed_columns"]:
        # ignore datefield > checked with datefields_list
        if c["key"]=="col_date":
            continue

        if c["columnname"] in cols:
            # column in csv found
            column_index[c["key"]] = cols.index(c["columnname"])
        elif "alt" in c and c["alt"] in cols:
            # column has 'alt'-entry in config and 'alt' is found in csv
            column_index[c["key"]] = cols.index(c["alt"])
        else:
            # column and 'alt' in csv not found
            raise ColumnNotFoundException(c["columnname"])

    # check for datefields
    counter = 0
    for d in datefields_list:
        d = d.lower()
        cols = list(map(lambda c: c.lower(), cols))
        if d in cols:
            column_index["col_date"+str(counter)] = cols.index(d)
            counter+=1
        else:
            raise ColumnNotFoundException(d)

def convert_line(line, linenumber):
    result = []
    # remove double quotes (generated by Griffeye in case of quot in fieldcontent)
    line = line.replace("\"\"", "'")

    while line.find('"') > -1:
        # cut out field with separator in it
        pos_start = line.find('"')+1
        while True:
            pos_end = line.find('"', pos_start)
            if pos_start == pos_end:
                pos_start += 1
                continue
            break
        if pos_start == 0 or pos_end == -1:
            raise LineNotValidException(linenumber)
        field = line[pos_start:pos_end]
        # get range before and after the field
        second_part = line[pos_end+2:-1]   # +1 = ", +1 = separator, -1 = general
        if pos_start > 2:
            first_part = line[:pos_start-2]  # -1 = ", -1 = general 
            # add first part to result_list
            result = result + first_part.split(csv_separator) + [field]
        else:
            # add field to result_list
            result = result + [field]

        # cut first part of line
        line = second_part
    result = result + second_part.split(csv_separator)
    return result

def analyze_header(filename):
    """
    - check for needed columns & fill columnindex-dictionary
    - sets the column count
    """
    with open(filename, "r", encoding=input_encoding) as file_input:
        header = file_input.readline().strip('\n')
        header = header.replace("\ufeff", "")
        if not args.s:
            detect_separator(header)
        check_columns(header)
        global column_count
        column_count = header.count(csv_separator)

def process_file():
    file_input = open(input_filename, "r", encoding=input_encoding)
    counter = 0
    for line in file_input:
        exclude = False
        counter += 1
        if counter == 1:
            # ignore csv-header
            continue

        # get data from file
        try:
            if line.count(csv_separator) != column_count:
                column = convert_line(line, counter+1)
            else:
                column = line.split(csv_separator)

            date_obj = get_date_field(column)
            data_device = column[column_index['col_device']]
            # create device when needed
            if data_device not in devices.keys():
                devices[data_device] = Device(data_device)
            device = devices[data_device]
            data_path = column[column_index['col_path']]
            data_type = column[column_index['col_type']]
            data_category = column[column_index['col_category']]
            data_hash = column[column_index['col_hash']]
            # cancel if path contains exclude text
            for e in exclude_list:
                if e.lower() in data_path.lower():
                    exclude = True
                    break
            if exclude:
                continue
            # separate thumbcaches from "normal" paths if its a thumb
            if not include_thumbcache and is_thumbcache(data_path):
                device.add_separate_thumb(data_category, data_path, data_type, data_hash)
                continue

            device.add_file(data_category, data_path, data_type, date_obj, data_hash)
        except LineNotValidException as exp:
            invalid_lines.append(exp.args[0])

        # update progressbar
        progress(counter, line_count)
    file_input.close()
    return counter

def calculate_device_totals():
    global cat_totals
    global cat_devcount
    for d in devices:
        categories = devices[d].get_categories()
        for dev_cat in categories.values():
            # increase/generate devicecount for category
            if dev_cat.name not in cat_devcount.keys():
                cat_devcount[dev_cat.name] = 1
            else:
                cat_devcount[dev_cat.name] += 1
            # get/generate total category
            total_cat = None
            if dev_cat.name not in cat_totals.keys():
                total_cat = Category(dev_cat.name)
                cat_totals[dev_cat.name] = total_cat
            else:
                total_cat = cat_totals[dev_cat.name]
            # merge device-category to total-category
            total_cat.merge(dev_cat)

def write_outputfile_docx():
    text_fontname = "Arial"
    text_fontsize = Pt(11)
    table_fontsize = Pt(8)
    table_rowheight = Pt(14)
    table_colwidth = Pt(160)
    table_2ndcol = Pt(280)

    document = Document()
    # write results of file-analysis
    document.add_heading(f"GRIFFEYE-CRAWLER - {labels['result_from']} {datetime.now().strftime('%d.%m.%Y')}", 1)
    p = document.add_paragraph()
    run = p.add_run(f"{labels['analyzed_file']}\t{input_filename}\n{labels['number_of_rows']}\t{line_count}\n{labels['defined_datefields']}\t{', '.join(datefields_list)}\n{labels['defined_excludes']}\t{', '.join(exclude_list)}\n{labels['thumbcaches_included']}\t{include_thumbcache}\n")
    run.font.name = text_fontname
    run.font.size = text_fontsize
    counter = 0
    totallength = len(devices)+1 # + total-table

    # write total results
    document.add_heading(f"{labels['total_over_all_devices']}", 2)
    for c in sorted(category_sort.keys()):
        if category_sort[c] not in cat_totals.keys():
            continue
        cat = cat_totals[category_sort[c]]
        # write table...
        table = document.add_table(rows=1, cols=2, style="Table Grid")
        # format header
        hdr_cells = table.rows[0].cells
        # cell merging
        hdr_cells[0].text = cat.name
        hdr_cells[0].width = table_colwidth
        datentr = labels['on_1_disk']
        if cat_devcount[category_sort[c]] > 1:
            datentr = labels['on_x_disks']
        hdr_cells[1].text = f"{cat.get_counts_string()} {labels['x_on_x']} {cat_devcount[category_sort[c]]} {datentr}"
        hdr_cells[1].width = table_2ndcol
        # background color
        cellshade = OxmlElement("w:shd")
        cellshade.set(qn("w:fill"), "#CCCCCC")
        cellprop = hdr_cells[0]._tc.get_or_add_tcPr()
        cellprop.append(cellshade)
        cellshade = OxmlElement("w:shd")
        cellshade.set(qn("w:fill"), "#CCCCCC")
        cellprop = hdr_cells[1]._tc.get_or_add_tcPr()
        cellprop.append(cellshade)
        # row alignment
        table.rows[0].height = table_rowheight
        hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # font
        run = hdr_cells[0].paragraphs[0].runs[0]
        run.font.name = text_fontname
        run.font.size = table_fontsize
        run.font.bold = True
        run = hdr_cells[1].paragraphs[0].runs[0]
        run.font.name = text_fontname
        run.font.size = table_fontsize

        # fill data...
        if cat.visible:
            # timeline
            row_cells = table.add_row().cells
            row_cells[0].text = f"{labels['distribution_in_time_period']}"
            row_cells[1].text = f"{cat.get_grouped_years()}"
            # proportion storage <-> browser cache
            row_cells = table.add_row().cells
            row_cells[0].text = f"{labels['percentage_browsercache']}"
            row_cells[1].text = f"{get_browser_percent(cat.get_browsercache_total(), cat.get_counts()[0])}"
            # show separated thumbcaches
            if not include_thumbcache:
                row_cells = table.add_row().cells
                row_cells[0].text = f"{labels['thumbcaches']}"
                row_cells[1].text = f"{cat.get_separate_thumbs_total()} ({cat.get_separate_thumbs_total_unique()})"
                cellshade = OxmlElement("w:shd")
                cellshade.set(qn("w:fill"), "#CCCCCC")
                cellprop = row_cells[1]._tc.get_or_add_tcPr()
                cellprop.append(cellshade)
        
        # format table
        for row in table.rows[1:]:
            i=-1
            row.height = table_rowheight
            row.cells[0].width = table_colwidth
            row.cells[1].width = table_2ndcol
            for cell in row.cells:
                i+=1
                # cell alignment
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # font
                run = cell.paragraphs[0].runs[0]
                run.font.name = text_fontname
                run.font.size = table_fontsize
                if i == 0:
                    # name-column bold
                    run.font.bold = True
        document.add_paragraph().paragraph_format.space_after = Pt(0)

    counter += 1
    # update progressbar
    progress(counter, totallength)

    # write results of devices
    for d in devices:
        counter += 1
        document.add_heading(d, 2)
        for c in sorted(category_sort.keys()):
            if category_sort[c] not in devices[d].categories:
                continue

            cat = devices[d].get_category(category_sort[c])
            # write table...
            table = document.add_table(rows=1, cols=2, style="Table Grid")
            # format header
            hdr_cells = table.rows[0].cells
            # cell merging
            hdr_cells[0].merge(hdr_cells[1])
            hdr_cells[0].text = cat.name
            # background color
            cellprop = hdr_cells[0]._tc.get_or_add_tcPr()
            cellshade = OxmlElement("w:shd")
            cellshade.set(qn("w:fill"), "#CCCCCC")
            cellprop.append(cellshade)
            # row alignment
            table.rows[0].height = table_rowheight
            hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # font
            run = hdr_cells[0].paragraphs[0].runs[0]
            run.font.name = text_fontname
            run.font.size = table_fontsize
            run.font.bold = True

            # fill data...
            # count & mediatype
            row_cells = table.add_row().cells
            row_cells[0].text = f"{labels['quantity_filetype']}"
            row_cells[1].text = cat.get_counts_string()
            if cat.visible:
                # daterange
                row_cells = table.add_row().cells
                row_cells[0].text = f"{labels['creation_on_disk']}"
                row_cells[1].text = f"{cat.get_date_range_string()}"
                # timeline
                row_cells = table.add_row().cells
                row_cells[0].text = f"{labels['distribution_in_time_period']}"
                row_cells[1].text = f"{cat.get_grouped_years()}"
                # proportion storage <-> browser cache
                row_cells = table.add_row().cells
                row_cells[0].text = f"{labels['percentage_browsercache']}"
                row_cells[1].text = f"{get_browser_percent(cat.get_browsercache_total(), cat.get_counts()[0])}"
                # paths
                row_cells = table.add_row().cells
                # cell merging
                row_cells[0].merge(row_cells[1])
                # show top-paths
                title_paragraph = row_cells[0].paragraphs[0].add_run(f"{labels['most_common_locations']}")
                title_paragraph.font.name = text_fontname
                title_paragraph.font.size = table_fontsize
                title_paragraph.font.bold = True
                i = 0
                # copy the pathlist and add a thumbcache- and browsercache-entries with the total sums to the temporary copy
                temppaths = dict(cat.paths)
                if cat.get_thumbcache_sum() > 0:
                    temppaths[name_for_thumbcache] = cat.get_thumbcache_obj()
                browser_sums = cat.get_browsercache_sums()
                for b in browser_sums.keys():
                    temppaths[name_for_browsercache+" "+b] = browser_sums[b]

                # work with the temporary pathlist incl. the thumbcache-entry
                for k in sorted(temppaths, key=lambda name: temppaths[name].count_total, reverse=True):
                    i += 1
                    if i > number_of_showed_paths:
                        break
                    row_paragraph = row_cells[0].paragraphs[0].add_run(f"\n- {shorten_path(k)}")
                    row_paragraph.font.name = text_fontname
                    row_paragraph.font.size = table_fontsize
                    row_paragraph.font.bold = False
                if len(temppaths) == 0:
                    row_paragraph = row_cells[0].paragraphs[0].add_run(f"\n-")
                    row_paragraph.font.name = text_fontname
                    row_paragraph.font.size = table_fontsize
                    row_paragraph.font.bold = False

                 # show separated thumbcaches
                if not include_thumbcache:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"{labels['thumbcaches']}"
                    row_cells[1].text = f"{cat.get_separate_thumbs_total()} ({cat.get_separate_thumbs_total_unique()})"
            
            # format table
            r = 1
            for row in table.rows[1:]:
                r+=1
                c=-1
                row.height = table_rowheight
                for cell in row.cells:
                    c+=1
                    # cell alignment
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # font
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = text_fontname
                    run.font.size = table_fontsize
                    if r != 6 and c == 0: # ignore 'most common locations'
                        # name-column bold
                        run.font.bold = True
                        cell.width = table_colwidth
                    if r != 6 and c == 1: # ignore 'most common locations'
                        cell.width = table_2ndcol
            document.add_paragraph().paragraph_format.space_after = Pt(0)

        progress(counter, totallength)
        document.save(result_filename)

def write_outputfile_json():
    json_obj = { 
        "meta": { 
            "processing_date": datetime.now().strftime('%d.%m.%Y'),
            "analyzed_file": input_filename,
            "row_count": line_count,
            "defined_datefields": ', '.join(datefields_list),
            "defined_excludes": ', '.join(exclude_list),
            "thumbcaches_included": include_thumbcache
        }
    }

    counter = 0
    totallength = len(devices)+1 # + total-table

    # write total results
    json_obj["total_over_all_devices"] = []
    for c in sorted(category_sort.keys()):
        if category_sort[c] not in cat_totals.keys():
            continue
        cat = cat_totals[category_sort[c]]
        t_counts = cat.get_counts()
        u_counts = cat.get_unique_counts()
        dates = cat.get_date_range()
        tmp_obj = {
                "category": cat.name,
                "count_summary": cat.get_counts_string(),
                "picture_count": t_counts[1],
                "picture_count_unique": u_counts[1],
                "video_count": t_counts[2],
                "video_count_unique": u_counts[2],
                "device_count": cat_devcount[cat.name],
                "creation_summary": cat.get_date_range_string(),
                "creation_startdate": dates[0],
                "creation_enddate": dates[1],
                "distribution_over_time": cat.get_grouped_years(),
                "percentace_browsercache": get_browser_percent(cat.get_browsercache_total(), cat.get_counts()[0])
            }
        if not include_thumbcache:
            tmp_obj["separate_thumbcaches_summary"] = f"{cat.get_separate_thumbs_total()} ({cat.get_separate_thumbs_total_unique()})"
            tmp_obj["thumbcaches_count"] = cat.get_separate_thumbs_total()
            tmp_obj["thumbcaches_count_unique"] = cat.get_separate_thumbs_total_unique()
        json_obj["total_over_all_devices"].append(tmp_obj)
    # update progressbar with total
    counter += 1
    progress(counter, totallength)

    # write results of devices
    json_obj["per_device"] = []
    for d in devices:
        counter += 1
        dev_obj = { "device": d }
        dev_obj["categories"] = []
        for c in sorted(category_sort.keys()):
            if category_sort[c] not in devices[d].categories:
                continue
            cat = devices[d].get_category(category_sort[c])
            t_counts = cat.get_counts()
            u_counts = cat.get_unique_counts()
            dates = cat.get_date_range()

            # generat list of most common locations
            loc_list = []
            i = 0
            # copy the pathlist and add a thumbcache- and browsercache-entries with the total sums to the temporary copy
            temppaths = dict(cat.paths)
            if cat.get_thumbcache_sum() > 0:
                temppaths[name_for_thumbcache] = cat.get_thumbcache_obj()
            browser_sums = cat.get_browsercache_sums()
            for b in browser_sums.keys():
                temppaths[name_for_browsercache+" "+b] = browser_sums[b]
            # work with the temporary pathlist incl. the thumbcache-entry
            for k in sorted(temppaths, key=lambda name: temppaths[name].count_total, reverse=True):
                i += 1
                if i > number_of_showed_paths:
                    break
                loc_list.append(f"{shorten_path(k)}")

            # create device object
            tmp_obj = {
                    "category": cat.name,
                    "count_summary": cat.get_counts_string(),
                    "picture_count": t_counts[1],
                    "picture_count_unique": u_counts[1],
                    "video_count": t_counts[2],
                    "video_count_unique": u_counts[2],
                    "creation_summary": cat.get_date_range_string(),
                    "creation_startdate": dates[0],
                    "creation_enddate": dates[1],
                    "distribution_over_time": cat.get_grouped_years(),
                    "percentage_browsercache": get_browser_percent(cat.get_browsercache_total(), cat.get_counts()[0]),
                    "most_common_locations": loc_list
                }
            if not include_thumbcache:
                tmp_obj["separate_thumbcaches_summary"] = f"{cat.get_separate_thumbs_total()} ({cat.get_separate_thumbs_total_unique()})"
                tmp_obj["thumbcaches_count"] = cat.get_separate_thumbs_total()
                tmp_obj["thumbcaches_unique"] = cat.get_separate_thumbs_total_unique()
            dev_obj["categories"].append(tmp_obj)
        # add device object to list of devices
        json_obj["per_device"].append(dev_obj)
        # update progressbar
        progress(counter, totallength)

    # write to file
    json_file = open(result_filename, "w", encoding=result_encoding)
    json_file.write(json.dumps(json_obj, indent=2, ensure_ascii=False))
    json_file.close()

def write_outputfile_txt():
    file_result = open(result_filename,"w", encoding=result_encoding)
    # write results of file-analysis
    file_result.write(f"GRIFFEYE-CRAWLER - {labels['result_from']} {datetime.now().strftime('%d.%m.%Y')}\n")
    file_result.write("="*43+"\n")
    file_result.write(f"{labels['analyzed_file']}\t{input_filename}\n")
    file_result.write(f"{labels['number_of_rows']}\t{line_count}\n")
    file_result.write(f"{labels['defined_datefields']}\t{', '.join(datefields_list)}\n")
    file_result.write(f"{labels['defined_excludes']}\t{', '.join(exclude_list)}\n")
    file_result.write(f"{labels['thumbcaches_included']}\t{include_thumbcache}\n")
    file_result.write("\n")
    counter = 0
    totallength = len(devices)+1 # + total-table

    # write total results
    file_result.write("\n{}\n".format(get_titlestring(f"{labels['total_over_all_devices']}", "=")))
    for c in sorted(category_sort.keys()):
        if category_sort[c] not in cat_totals.keys():
            continue
        cat = cat_totals[category_sort[c]]
        file_result.write("\n{}\n".format(get_titlestring(cat.name, "\u0387")))
        # count & mediatype
        file_result.write(f"{labels['quantity_filetype']}\t\t\t\t{cat.get_counts_string()}\n")
        # devicecount
        file_result.write(f"{labels['number_of_devices']}\t\t\t\t{cat_devcount[cat.name]}\n")
        if cat.visible:
            # daterange
            file_result.write(f"{labels['creation_on_disk']}\t{cat.get_date_range_string()}\n")
            # timeline
            file_result.write(f"{labels['distribution_in_time_period']}\t{cat.get_grouped_years()}\n")
            # proportion storage <-> browser cache
            file_result.write(f"{labels['percentage_browsercache']}\t\t{get_browser_percent(cat.get_browsercache_total(), cat.get_counts()[0])}\n")
        # show separated thumbcaches
        if not include_thumbcache:
            file_result.write(f"{labels['thumbcaches']}\t\t\t{cat.get_separate_thumbs_total()} ({cat.get_separate_thumbs_total_unique()})\n")
    file_result.write("\n")

    counter += 1
    # update progressbar
    progress(counter, totallength)

    # write results of devices
    for d in devices:
        counter += 1
        file_result.write("\n{}\n".format(get_titlestring(d, "=")))
        for c in sorted(category_sort.keys()):
            if category_sort[c] not in devices[d].categories:
                continue

            cat = devices[d].get_category(category_sort[c])
            file_result.write("\n{}\n".format(get_titlestring(cat.name, "\u0387")))
            # count & mediatype
            file_result.write(f"{labels['quantity_filetype']}\t\t\t\t{cat.get_counts_string()}\n")
            if cat.visible:
                # daterange
                file_result.write(f"{labels['creation_on_disk']}\t\t\t\t{cat.get_date_range_string()}\n")
                # timeline
                file_result.write(f"{labels['distribution_in_time_period']}\t{cat.get_grouped_years()}\n")
                # proportion storage <-> browser cache
                file_result.write(f"{labels['percentage_browsercache']}\t\t{get_browser_percent(cat.get_browsercache_total(), cat.get_counts()[0])}\n")
                # paths
                file_result.write(f"{labels['most_common_locations']}")
                if len(cat.paths) == 0:
                    file_result.write("\t-")
                file_result.write("\n")
                # show top-paths
                i = 0
                # copy the pathlist and add a thumbcache- and browsercache-entries with the total sums to the temporary copy
                temppaths = dict(cat.paths)
                if cat.get_thumbcache_sum() > 0:
                    temppaths[name_for_thumbcache] = cat.get_thumbcache_obj()
                browser_sums = cat.get_browsercache_sums()
                for b in browser_sums.keys():
                    temppaths[name_for_browsercache+" "+b] = browser_sums[b]

                # work with the temporary pathlist incl. the thumbcache-entry
                for k in sorted(temppaths, key=lambda name: temppaths[name].count_total, reverse=True):
                    i += 1
                    if i > number_of_showed_paths:
                        break
                    file_result.write(f"- {shorten_path(k)}\n")
                # show separated thumbcaches
                if not include_thumbcache:
                    file_result.write(f"{labels['thumbcaches']}\t\t\t{cat.get_separate_thumbs_total()} ({cat.get_separate_thumbs_total_unique()})\n")
        file_result.write("\n")
        # update progressbar
        progress(counter, totallength)

    file_result.close()

def write_pathdetails():
    """
    creates the outputfile (txt) with detailed information
    """
    details_name = config["result"]["pathdetails_name"]
    if not details_name.endswith(".txt"):
        details_name = details_name+".txt"
    details_name = f"{get_file_basename(result_filename)}_{details_name}"
    enc = config["result"]["pathdetails_encoding"]
    file_result = open(get_output_path(input_filename)+details_name,"w", encoding=enc)
    # write results of file-analyze
    file_result.write(f"GRIFFEYE-CRAWLER - {labels['path_details_from']} {datetime.now().strftime('%d.%m.%Y')}\n")
    file_result.write("="*47+"\n")
    file_result.write(f"{labels['analyzed_file']}\t{input_filename}\n")
    file_result.write(f"{labels['number_of_rows']}\t{line_count}\n")
    file_result.write(f"{labels['defined_datefields']}\t{', '.join(datefields_list)}\n")
    file_result.write(f"{labels['defined_excludes']}\t{', '.join(exclude_list)}\n")
    file_result.write(f"{labels['thumbcaches_included']}\t{include_thumbcache}\n")
    file_result.write("\n")

    # write results of devices
    counter = 0
    for d in devices:
        counter += 1
        file_result.write("\n{}\n".format(get_titlestring(d, "=")))
        file_result.write(f"{devices[d].get_counts()[0]} {labels['files']} ({labels['legal']}: {devices[d].get_counts()[1]}, {labels['illegal']}: {devices[d].get_counts()[2]})")
        if (devices[d].get_counts()[0]==0):
            file_result.write("  >>  0%")
        else:
            file_result.write("  >>  {:.2f}% {}\n".format((devices[d].get_counts()[2]/devices[d].get_counts()[0])*100, labels['illegal']))
        for c in sorted(category_sort.keys()):
            if category_sort[c] not in devices[d].categories:
                continue

            cat = devices[d].get_category(category_sort[c])
            file_result.write("\n{}\n".format(get_titlestring(cat.name, "\u0387")))
            # count & mediatype
            file_result.write(f"{labels['quantity_filetype']}\t\t\t\t{cat.get_counts_string()}\n")
            # daterange
            file_result.write(f"{labels['creation_on_disk']}\t\t\t\t{cat.get_date_range_string()}\n")
            # timeline
            file_result.write(f"{labels['distribution_in_time_period']}\t{cat.get_grouped_years()}\n")
            # proportion storage <-> browser cache
            browser_total = cat.get_browsercache_total()
            counts_total = cat.get_counts()[0]
            perc = (browser_total/counts_total)*100 if counts_total > 1 else 0
            perc_str = "{:.0f}%".format(perc)
            if round(perc, 0) == 0 and perc > 0:
                perc_str = "<1%"
            file_result.write(f"{labels['percentage_browsercache']}\t\t{perc_str} >>> ({labels['total']}: {counts_total}, {labels['browsercache']}: {browser_total})\n")
            # paths
            file_result.write(f"{labels['locations']}\n")
            # show paths
            # copy the pathlist and add a thumbcache-entry with the total sum to the temporary copy
            temppaths = dict(cat.paths)
            if cat.get_thumbcache_sum() > 0:
                temppaths[name_for_thumbcache] = cat.get_thumbcache_obj()
            # work with the temporary pathlist incl. the thumbcache-entry
            for k in sorted(temppaths, key=lambda name: temppaths[name].count_total, reverse=True):
                path = temppaths[k]
                details_text = f" (p: {path.count_pic}, v: {path.count_vid})" if path.show_details else ""
                file_result.write(f"- {k} >>> {path.count_total} {details_text}\n")
            # separated thumbcaches
            if not include_thumbcache:
                file_result.write(f"{labels['thumbcaches']}\t\t\t{cat.get_separate_thumbs_total()} ({cat.get_separate_thumbs_total_unique()})\n")
                for p in sorted(cat.separate_thumbs, key=lambda path: cat.separate_thumbs[path].count_total, reverse=True):
                    path = cat.separate_thumbs[p]
                    details_text = f" (p: {path.count_pic}, v: {path.count_vid})" if path.show_details else ""
                    file_result.write(f"- {p} >>> {path.count_total} {details_text}\n")
            # if available, write other caches
            if len(cat.caches)>0:
                file_result.write(f"    > {labels['caches']} <\n")
                sorted_caches = []
                for cache in sorted(cat.caches.values(), key=lambda c: cat.caches[c.name].count, reverse=True):
                    sorted_caches.append(cache)

                for cache in sorted_caches:
                    file_result.write(f"- {cache.name} >>> {cache.count}\n")
                file_result.write(f"    > {labels['cache_details']} <\n")
                for cache in sorted_caches:
                    file_result.write(f"{cache.name}\n")
                    for path in sorted(cache.paths.values(), key=lambda p: p.count_total, reverse=True):
                        details_text = f" (p: {path.count_pic}, v: {path.count_vid})" if path.show_details else ""
                        file_result.write(f"- {path.path} >>> {path.count_total} {details_text}\n")

        file_result.write("\n")
        # update progressbar
        progress(counter, len(devices))

    file_result.close()

def read_config():
    """
    read configurations from config.json (can be overwritten by input options)
    """
    global config
    global input_encoding
    global result_encoding
    global result_language
    global category_legality
    global category_visibilty
    global category_sort
    global known_cache_paths
    global known_cache_names
    global number_of_showed_paths
    global include_thumbcache
    global date_format

    with open('config.json', 'r', encoding='utf-8') as c:
        data = c.read()
    config = json.loads(data)
    input_encoding = config["input"]["encoding"]
    result_encoding = config["result"]["encoding"]
    result_language = config["result"]["language"]
    for cat in config["categories"]:
        category_legality[cat["name"]] = cat["legality"]
        category_visibilty[cat["name"]] = cat["show_in_report"]
        category_sort[cat["sort"]] = cat["name"]
    for cac in config["caches"]:
        name = cac["name"]
        known_cache_paths[cac["path"]] = name
        if name not in known_cache_names:
            known_cache_names[name] = CacheGroup(name, cac["is_browser"], cac["is_thumbcache"])
        group = known_cache_names[name]
        group.add_pattern(cac["path"])

    number_of_showed_paths = config["result"]["number_of_showed_paths"]
    include_thumbcache = config["other"]["include_thumbcache"]
    date_format = config["input"]["date_format"]

def read_labels():
    """ read labels for multi language support from labels.json """
    global labels
    global result_language
    done = False

    with open('labels.json', 'r', encoding='utf-8') as d:
        data = d.read()
    config = json.loads(data)
    for l in config["languages"]:
        if l["lang"] == result_language:
            for lab in l["labels"]:
                labels[lab["label"]] = lab["text"]
            result_language = l["lang"]
            done = True
            break
    
    if not done:
        # no language found > check for main language
        for l in config["languages"]:
            if l["lang"][:2] == result_language[:2]:
                for lab in l["labels"]:
                    labels[lab["label"]] = lab["text"]
                print(f"[i] Language '{result_language}' not found! '{result_language[:2]}' used instead...")
                result_language = l["lang"][:2]
                done = True
                break
    
    if not done:
        # no language found > throw error
        raise LanguageNotFoundException(result_language)
    

def generate_datefields_list():
    global datefields_list
    if args.date:
        datefields_list = args.date.split(",")
            
    if not datefields_list:
        if args.date:
            print(f"[i] No date definitions found! Default is used...")
        for c in config["needed_columns"]:
            if c["key"]=="col_date":
                datefields_list.append(c["columnname"])
                break
        datefields_list.append(config["other"]["alternative_date_column"])

def generate_exclude_list():
    global exclude_list
    if args.exclude:
        exclude_list = args.exclude.split(",")

def has_file_extension(input):
    return os.path.splitext(input)[1]!=""

def get_file_basename(input):
    filename = os.path.basename(input)
    return os.path.splitext(filename)[0]

def get_output_format():
    ext = ""
    if args.o and has_file_extension(args.o):
        ext = os.path.splitext(args.o)[1][1:] # remove . at start of extension
    elif args.f:
        ext = args.f
    else:
        ext = default_format
    
    # check extension
    if ext in valid_formats:
        return ext
    
    print(f"[i] Output format '{ext}' not found! Default format is used...")
    return default_format

def get_output_name(inputname):
    if args.o and has_file_extension(args.o):
        return f"{get_file_basename(args.o)}.{result_format}"
    return f"{get_file_basename(inputname)}.{result_format}"

def get_output_path(inputname):
    path = ""
    if args.o:
        if not has_file_extension(args.o):
            path = args.o
        else:
            path = os.path.dirname(args.o)
    else:
        path = os.path.dirname(inputname)

    # check for existance
    if path != "" and not os.path.exists(path):
        raise PathNotFoundException(path)
    if path != "":
        path = path+os.sep
    return path



# init
column_index = {}
devices = {}
cat_totals = {}
cat_devcount = {}
invalid_lines = []
datefields_list = []
exclude_list = []
csv_separator = ""
column_count = 0
line_count = 0

default_format = "docx"
valid_formats = ["docx", "json", "txt"]

empty_date = datetime.strptime("01.01.0001", "%d.%m.%Y")
unix_date = datetime.strptime("01.01.1970", "%d.%m.%Y")

# init configs
config = {}
input_encoding = ""
result_filename = ""
result_encoding = ""
result_language = "en"
labels = {}
category_legality = {}
category_visibilty = {}
category_sort = {}
known_cache_paths = {}
known_cache_names = {}
number_of_showed_paths = 0
include_thumbcache = False
date_format = ""

# init argparse
args = None
configure_argparse()

try:
    read_config()
    input_filename = args.file
    # remove " & ' from path (prevents error while reading the file)
    input_filename = input_filename.replace("\"", "")
    input_filename = input_filename.replace("'", "")

    result_format = get_output_format()
    result_filename = os.path.join(get_output_path(input_filename), get_output_name(input_filename))
    # get linecount for progressbar
    line_count = get_linecount(input_filename)
    # set separator from options (deactivates automatic detection)
    csv_separator = args.s if args.s else csv_separator
    # set dateformat from options
    date_format = args.d if args.d else date_format
    # set number of showed paths from options
    number_of_showed_paths = args.n if args.n else number_of_showed_paths
    # set number of showed paths from options
    include_thumbcache = args.includethumbs if args.includethumbs else include_thumbcache
    # set language from options
    result_language = args.l if args.l else result_language
    read_labels()
    # set list of datefields
    generate_datefields_list()
    # set list of excludes
    generate_exclude_list()

    # analyze file
    analyze_header(input_filename)
    # process data
    print(f"Processing records in '{input_filename}'...")
    processed = process_file()
    if len(invalid_lines) > 0:
        print()
        print("  [i] Invalid rows detected in CSV and ignored in processing")
        print("  [i] Rows: ", end="")
        for l in invalid_lines:
            print(l, end="  ")
        print()
    print()
    calculate_device_totals()

    # write output-files
    print("Write result files...")
    name_for_thumbcache = config["other"]["name_for_thumbcache"]
    name_for_browsercache = config["other"]["name_for_browsercache"]
    if result_format == "txt":
        write_outputfile_txt()
    elif result_format == "json":
        write_outputfile_json()
    elif result_format == "docx":
        write_outputfile_docx()

    if config["result"]["generate_pathdetails"] and not args.nodetails:
        write_pathdetails()

    print()
    print()
    print(f"DONE! {processed} record processed (check results in '{result_filename}')")

except PathNotFoundException as exp:
    print()
    print("[!] Processing aborted!")
    print(">", exp.message)
except ColumnNotFoundException as exp:
    print()
    print("[!] Processing aborted!")
    print(">", exp.message)
except SeparatorNotFoundException as exp:
    print()
    print("[!] Processing aborted!")
    print(">", exp.message)
except LanguageNotFoundException as exp:
    print()
    print("[!] Processing aborted!")
    print(">", exp.message)
except FileNotFoundError as exp:
    print()
    print("[!] Processing aborted!")
    print(f"> File '{exp.filename}' not found")
except KeyError as exp:
    print()
    print("[!] Processing aborted!")
    print(f"> Configuration '{exp}' not found")
except UnicodeDecodeError as exp:
    print()
    print("[!] Processing aborted!")
    if exp.args[0] == "utf-8":
        print("File is not in UTF-8 format. Please adjust configuration or convert the file...")
    else:
        print("File is in an unknown format")
except UnicodeError as exp:
    print()
    print("[!] Processing aborted!")
    if "UTF-16" in exp.args[0]:
        print("File is not in UTF-16 format. Please adjust configuration or convert the file...")
    else:
        print("File is in an unknown format")
except Exception as exp:
    print()
    print("[!] Processing aborted!")
    traceback.print_exc()

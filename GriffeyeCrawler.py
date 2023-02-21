#!/usr/bin/env python
#-*- coding:utf-8 -*-
"""
GRIFFEYE-CRAWLER
----------------
Analysiert eine exportierte Dateiliste aus Griffeye pro Gerät & Kategorie
- Summiert die Bilder und Videos (Total & binary unique)
- Fasst die Dateipfade zusammen und unterteilt diese in Cache- & Nicht-Cache-Pfade auf
- Ermittelt die Pfade mit den meisten Inhalten
- Ermittelt den gesamten Zeitraum der Dateierstellung
- Ermittelt die prozentuelle Verteilung der Dateierstellung im betroffenen Zeitraum
- Ermittelt das prozentuelle Verhältnis im Browsercache und der übrigen Ablage
- Ermittelt sämtliche erwähnten Punkte auch als Total über alle Geräte + die Anzahl der betroffenen Geräte
- Generiert eine Ergebnisdatei im TXT oder DOCX-Format

(c) 2023, Luzerner Polizei
Author:  Michael Wicki
Version: 0.5.2
"""
version = "v0.5.2"

import os
import sys
import json
import traceback
from datetime import datetime
# docx...
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL



class Device:
    def __init__(self, sourceid, category, initial_date):
        self.sourceid = sourceid
        self.categories = { category: Category(category, initial_date) }
        self.legal_count = 0
        self.illegal_count = 0

    def addDate(self, category, date):
        if category not in self.categories.keys():
            self.categories[category] = Category(category, date)
        else:
            self.categories[category].addDate(date)

    def addFile(self, category, path, mediatype, date, hash):
        self.categories[category].addFile(path, mediatype, date, hash)
        
        # increase legal/illegal count
        if category_legality.get(category, True):
            self.legal_count += 1
        else:
            self.illegal_count += 1

    def getSourceId(self):
        return self.sourceid

    def getCategories(self):
        return self.categories

    def getCategory(self, category):
        if category not in self.categories.keys():
            return None
        else:
            return self.categories[category]

    def getCounts(self):
        """
        returns a tuple with total count, legal count & illegal count of the device
        """
        return (self.legal_count+self.illegal_count, self.legal_count, self.illegal_count)


class Category:
    def __init__(self, name, initial_date):
        self.name = name
        self.legality = category_legality.get(name, True)
        self.min_date = empty_date
        if initial_date != unix_date:
            self.min_date = initial_date
        self.max_date = initial_date
        self.date_groups = {}
        self.pic_count = 0
        self.vid_count = 0
        self.tot_count = 0
        self.paths = {}
        self.cachepaths = {}
        self.cachegroups = {}
        self.pic_hashes = set()
        self.vid_hashes = set()

    def addDate(self, date):
        if date != empty_date and date != unix_date:
            if self.min_date == empty_date or date < self.min_date:
                self.min_date = date
            if self.max_date == empty_date or date > self.max_date:
                self.max_date = date

    def addFile(self, path, mediatype, date, hash):
        # increase counters & add hash to 'hashes' (>> deduplicates itself)
        self.tot_count += 1
        if mediatype == "Image":
            self.pic_count += 1
            self.pic_hashes.add(hash)
        if mediatype == "Video":
            self.vid_count += 1
            self.vid_hashes.add(hash)

        # increase path
        self.increasePath(path)

        # increase date
        self.increaseDate(date)

    def merge(self, merge_cat):
        # merge daterange
        self.addDate(merge_cat.min_date)
        self.addDate(merge_cat.max_date)
        # merge counts
        self.tot_count += merge_cat.getCounts()[0]
        self.pic_count += merge_cat.getCounts()[1]
        self.vid_count += merge_cat.getCounts()[2]
        # merge dategroups
        for item in merge_cat.date_groups.items():
            key = item[0]
            value = item[1]
            if key not in self.date_groups.keys():
                self.date_groups[key] = value  # create
            else:
                self.date_groups[key] += value # increase
        # merge paths
        for path in merge_cat.paths.keys():
            if path not in self.paths.keys():
                self.paths[path] = 1    # create
            else:
                self.paths[path] += 1   # increase
        # merge cachepaths
        for path in merge_cat.cachepaths.keys():
            group = self.getCacheGroup(path)
            if group is not None:
                if group not in self.cachegroups.keys():
                    self.cachegroups[group] = merge_cat.cachepaths[path]    # create
                else:
                    self.cachegroups[group] += merge_cat.cachepaths[path]     # increase
        # merge hashes
        self.pic_hashes.update(merge_cat.getPicHashset())
        self.vid_hashes.update(merge_cat.getVidHashset())

    def getCacheGroup(self, path):
        for k in known_cache_paths.keys():
            if k in path:
                return known_cache_paths[k]
        return None
    
    def increasePath(self, path):
        if path in self.cachepaths.keys():
            # path is in cachepath >> increase count
            self.cachepaths[path] += 1
            self.cachegroups[self.getCacheGroup(path)] += 1
        else:
            # path NOT in cachepath >> check for cache
            group = self.getCacheGroup(path)
            if group is not None:
                # path is cache
                self.cachepaths[path] = 1 # create
                if group not in self.cachegroups.keys():
                    self.cachegroups[group] = 1    # create
                else:
                    self.cachegroups[group] += 1   # increase
            else:
                # path is NOT cache
                if path not in self.paths.keys():
                    self.paths[path] = 1    # create
                else:
                    self.paths[path] += 1   # increase

    def increaseDate(self, date):
        year = int(date[6:10])
        if year == 1 or year == 1970: # no date or unix date
            year = 9999
        if year not in self.date_groups.keys():
            self.date_groups[year] = 1  # create
        else:
            self.date_groups[year] += 1 # increase

    def getDateRange(self):
        if self.min_date == empty_date or self.max_date == empty_date:
            return "undefiniert"
        return self.min_date.strftime("%d.%m.%Y")+" - "+self.max_date.strftime("%d.%m.%Y")

    def getDateRangeDays(self):
        return (self.max_date - self.min_date).days + 1

    def getPicHashset(self):
        return self.pic_hashes

    def getVidHashset(self):
        return self.vid_hashes

    def getUniqueCounts(self):
        """
        returns a tuple with total count, picture count & video count of binary unique files (based on the hash)
        """
        return (len(self.pic_hashes.union(self.vid_hashes)), len(self.pic_hashes), len(self.vid_hashes))

    def getCounts(self):
        """
        returns a tuple with total count, picture count & video count of the category
        """
        return (self.tot_count, self.pic_count, self.vid_count)
    
    def getCountsString(self):
        """
        returns a string with formatted picture- & videos-count
        """
        result = ""
        # pictures
        if self.pic_count > 0:
            result += "{} ".format(self.pic_count)
            if self.pic_count > 1:
                result += "Bilder"
            else:
                result += "Bild"
            # binary unique
            result += " ({})".format(len(self.pic_hashes))
            if self.vid_count > 0:
                result += ", "
        # videos
        if self.vid_count > 0:
            result += "{} ".format(self.vid_count)
            if self.vid_count > 1:
                result += "Videos"
            else:
                result += "Video"
            # binary unique
            result += " ({})".format(len(self.vid_hashes))
        return result

    def getGroupedDates(self):
        """
        returns a string with the percentage of illegal files per year
        """
        result = ""
        for year in sorted(self.date_groups.keys()):
            # calculate percentage of total files
            perc = (self.date_groups[year]/self.tot_count)*100
            if year == 9999:
                year = "undef."
            perc_str = "{:.0f}%".format(perc)
            if round(perc, 0) == 0 and perc > 0:
                perc_str = "<1%"
            result = result+"{}: {}, ".format(year, perc_str)
        return result[:-2] # kill last ', '

    def getBrowserCacheTotal(self):
        sum = 0
        for c in self.cachegroups.keys():
            if c in browser_names:
                sum += self.cachegroups[c]
        return sum

    def getBrowserCacheSums(self):
        """
        returns a dict with counts (value) for the specific browsers (key)
        """
        result = {}
        for c in self.cachegroups.keys():
            if c in browser_names:
                result[c] = self.cachegroups[c]
        return result

    def getThumbcacheSum(self):
        sum = 0
        for c in self.cachegroups.keys():
            if c in thumbcache_names:
                sum += self.cachegroups[c]
        return sum


class ColumnNotFoundException(Exception):
    """
    error in case of a column name not found
    """
    def __init__(self, columnname):
        self.message = "Spalte '{}' wurde nicht gefunden".format(columnname)

class LineNotValidException(Exception):
    """
    error in case of a csv-entry with ; in a field without " around it
    """
    def __init__(self, linenumber):
        self.message = "Zeile '{}' ist ungültig".format(linenumber)

class ResultFormatUnknownException(Exception):
    """
    error in case of an undefined output format
    """
    def __init__(self, format):
        self.message = "Ausgabeformat unbekannt ('{}')".format(format)


def progress(count, total, status=''):
    """
    handling of the progressbar
    """
    bar_len = 60
    filled_len = int(round(bar_len * count / float(total)))
    percents = round(100.0 * count / float(total), 1)
    bar = '#' * filled_len + '-' * (bar_len - filled_len)

    sys.stdout.write('[%s] %s%s ...%s\r' % (bar, percents, '%', status))
    sys.stdout.flush()

def getLinecount(filename):
    """
    count total lines for progressbars
    """
    counter = -1
    file_input = open(filename, "r", encoding=input_encoding)
    for line in file_input:
        counter += 1
        if counter == 0:
            continue
    file_input.close()
    return counter

def getTitleString(title, symbol="-", length=70):
    """
    creates a titleline with centered text
    """
    half_length = (length//2)-1  # including blank
    half_title = (len(title)//2)-1  # including blank
    symbol_count = half_length-half_title
    addition = ""
    if half_title%2 > 0 or half_length%2 > 0:
        addition = symbol
    return symbol*symbol_count+" "+title+" "+symbol*symbol_count+addition

def getBrowserPercent(browser_count, total_count):
    perc = (browser_count/total_count)*100
    if round(perc, 0) == 0 and perc > 0:
        return "<1%"
    return "{:.0f}%".format(perc)

def shortenPath(path):
    """
    shortens the filepath by the first two directories
    """
    first = path[path.find(os.path.sep)+1:]
    return first[first.find(os.path.sep)+1:]

def detectSeparator(header):
    """
    detect the csv separator (, or ;)
    """
    global csv_separator
    if header.find(',') > -1:
        csv_separator = ","
    elif header.find(';') > -1:
        csv_separator = ";"
    else:
        csv_separator = input("CSV-Separator konnte nicht ermittelt werden... Durch welches Zeichen werden die Spalten getrennt?")

def checkColumns(header):
    """
    check for needed columns & fill columnindex-dictionary for column access with columnname
    """
    cols = header[:-1].split(csv_separator)
    for c in config["needed_columns"]:
        if c["columnname"] in cols:
            # column in csv found
            column_index[c["key"]] = cols.index(c["columnname"])
        elif "alt" in c and c["alt"] in cols:
            # column has 'alt'-entry in config and 'alt' is found in csv
            column_index[c["key"]] = cols.index(c["alt"])
        else:
            # column and 'alt' in csv not found
            raise ColumnNotFoundException(c["columnname"])

    # check for alternative date for the case of empty date (01.01.0001) in mobiles
    alt_date = config["other"]["alternative_date_column"]
    alt_key = config["other"]["alternative_date_key"]
    column_index[alt_key] = -1
    if alt_date in cols:
        column_index[alt_key] = cols.index(alt_date)

def convertLine(line, linenumber):
    result = []
    while line.find('"') > -1:
        # cut out field with separator in it
        pos_start = line.find('"')+1
        while True:
            pos_end = line.find('"', pos_start)
            if pos_start == pos_end:
                pos_start += 1
                continue
            break
        if pos_start == 0 or pos_end == -1:
            raise LineNotValidException(linenumber)
        field = line[pos_start:pos_end]
        # get range before and after the field
        second_part = line[pos_end+2:-1]   # +1 = ", +1 = separator, -1 = general
        if pos_start > 2:
            first_part = line[:pos_start-2]  # -1 = ", -1 = general 
            # add first part to result_list
            result = result + first_part.split(csv_separator) + [field]
        else:
            # add field to result_list
            result = result + [field]

        # cut first part of line
        line = second_part
    result = result + second_part.split(csv_separator)
    return result

def analyzeFile(filename):
    """
    - check for needed columns & fill columnindex-dictionary
    - collect devices & fill devicelist
    - detect min & max date for daterange (per device)
    """
    counter = -1
    file_input = open(filename, "r", encoding=input_encoding)
    for line in file_input:
        counter += 1
        if counter == 0:
            # csv-header...
            line = line.replace("\ufeff", "")
            detectSeparator(line)
            checkColumns(line)
            global column_count
            column_count = line.count(csv_separator)
            continue

        # csv-entry...
        # get device & date from csv
        try:
            if line.count(csv_separator) != column_count:
                data = convertLine(line, counter+1)
            else:
                data = line.split(csv_separator)
            data_category = data[column_index['col_category']]
            data_date = data[column_index['col_date']]
            # if date is empty (01.01.0001) try the alternative date
            if data_date == empty_date_string and column_index['col_alt_date'] > -1:
                data_date = data[column_index['col_alt_date']]
            current_date = datetime.strptime(data_date[0:10], "%d.%m.%Y")
            data_device = data[column_index['col_device']]
            # check for device or create it when needed
            if data_device not in devices.keys():
                devices[data_device] = Device(data_device, data_category, current_date)
            else:
                devices[data_device].addDate(data_category, current_date)
        except LineNotValidException as exp:
            invalid_lines.append(exp.args[0])
        # update progressbar
        progress(counter, linecount)

    file_input.close()
    return counter

def writeOutputfileTxt():
    file_result = open(result_filename,"w", encoding=result_encoding)
    # write results of file-analysis
    file_result.write("GRIFFEYE-CRAWLER - Ergebnis vom {}\n".format(datetime.now().strftime("%d.%m.%Y")))
    file_result.write("="*43+"\n")
    file_result.write("Analysierte Datei:     {}\n".format(input_filename))
    file_result.write("Anzahl Datensätze:     {}\n".format(linecount))
    file_result.write("\n")
    counter = 0
    totallength = len(devices)+1 # + total-table

    # write total results
    file_result.write("\n{}\n".format(getTitleString("Total über alle Geräte", "=")))
    for c in sorted(category_sort.keys()):
        if category_sort[c] not in total.keys():
            continue
        cat = total[category_sort[c]]
        file_result.write("\n{}\n".format(getTitleString(cat.name, "\u0387")))
        # count & mediatype
        file_result.write("Menge/Dateityp:\t")
        file_result.write("{}\n".format(cat.getCountsString()))
        # devicecount
        file_result.write("Anzahl Datenträger:\t{}".format(cat_devcount[category_sort[c]]))
        file_result.write("\n")
        if category_sort[c] != "Legale Pornographie":
            # daterange
            file_result.write("Erstellung auf Datenträger:\t{}".format(cat.getDateRange()))
            file_result.write("\n")
            # timeline
            file_result.write("Verteilung im Zeitraum:\t{}".format(cat.getGroupedDates()))
            file_result.write("\n")
            # proportion storage <-> browser cache
            file_result.write("Anteil Browsercache:\t{}".format(getBrowserPercent(cat.getBrowserCacheTotal(), cat.getCounts()[0])))
            file_result.write("\n")
    file_result.write("\n")

    counter += 1
    # update progressbar
    progress(counter, totallength)

    # write results of devices
    for d in devices:
        counter += 1
        file_result.write("\n{}\n".format(getTitleString(d, "=")))
        for c in sorted(category_sort.keys()):
            if category_sort[c] not in devices[d].categories:
                continue

            cat = devices[d].getCategory(category_sort[c])
            file_result.write("\n{}\n".format(getTitleString(cat.name, "\u0387")))
            # count & mediatype
            file_result.write("Menge/Dateityp:\t")
            file_result.write("{}\n".format(cat.getCountsString()))
            if category_sort[c] != "Legale Pornographie":
                # daterange
                file_result.write("Erstellung auf Datenträger:\t{}".format(cat.getDateRange()))
                file_result.write("\n")
                # timeline
                file_result.write("Verteilung im Zeitraum:\t{}".format(cat.getGroupedDates()))
                file_result.write("\n")
                # proportion storage <-> browser cache
                file_result.write("Anteil Browsercache:\t{}".format(getBrowserPercent(cat.getBrowserCacheTotal(), cat.getCounts()[0])))
                file_result.write("\n")
                # paths
                file_result.write("Häufigste Speicherorte:\t")
                file_result.write("\n")
                # show top-paths
                i = 0
                # copy the pathlist and add a thumbcache- and browsercache-entries with the total sums to the temporary copy
                temppaths = dict(cat.paths)
                thumbsum = cat.getThumbcacheSum()
                if thumbsum > 0:
                    temppaths[name_for_thumbcache] = thumbsum
                browser_sums = cat.getBrowserCacheSums()
                for b in browser_sums.keys():
                    temppaths[name_for_browsercache+" "+b] = browser_sums[b]

                # work with the temporary pathlist incl. the thumbcache-entry
                for k in sorted(temppaths, key=temppaths.get, reverse=True):
                    i += 1
                    if i > config["result"]["number_of_showed_paths"]:
                        break
                    file_result.write("- {}\n".format(shortenPath(k)))

        file_result.write("\n")
        # update progressbar
        progress(counter, totallength)

    file_result.close()

def writeOutputfileDocx():
    text_fontname = "Arial"
    text_fontsize = Pt(11)
    table_fontsize = Pt(8)
    table_rowheight = Pt(14)

    document = Document()
    # write results of file-analysis
    document.add_heading("GRIFFEYE-CRAWLER - Ergebnis vom {}".format(datetime.now().strftime("%d.%m.%Y")), 1)
    p = document.add_paragraph()
    run = p.add_run("Analysierte Datei:\t{}\nAnzahl Datensätze:\t{}".format(input_filename, linecount))
    run.font.name = text_fontname
    run.font.size = text_fontsize
    counter = 0
    totallength = len(devices)+1 # + total-table

    # write total results
    document.add_heading("Total über alle Geräte", 2)
    for c in sorted(category_sort.keys()):
        if category_sort[c] not in total.keys():
            continue
        cat = total[category_sort[c]]
        # write table...
        table = document.add_table(rows=1, cols=2, style="Table Grid")
        # format header
        hdr_cells = table.rows[0].cells
        # cell merging
        hdr_cells[0].text = cat.name
        datentr = "Datenträger"
        if cat_devcount[category_sort[c]] > 1:
            datentr = "Datenträgern"
        hdr_cells[1].text = "{} auf {} {}".format(cat.getCountsString(), cat_devcount[category_sort[c]], datentr)
        # background color
        cellshade = OxmlElement("w:shd")
        cellshade.set(qn("w:fill"), "#CCCCCC")
        cellprop = hdr_cells[0]._tc.get_or_add_tcPr()
        cellprop.append(cellshade)
        cellshade = OxmlElement("w:shd")
        cellshade.set(qn("w:fill"), "#CCCCCC")
        cellprop = hdr_cells[1]._tc.get_or_add_tcPr()
        cellprop.append(cellshade)
        # row alignment
        table.rows[0].height = table_rowheight
        hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # font
        run = hdr_cells[0].paragraphs[0].runs[0]
        run.font.name = text_fontname
        run.font.size = table_fontsize
        run.font.bold = True
        run = hdr_cells[1].paragraphs[0].runs[0]
        run.font.name = text_fontname
        run.font.size = table_fontsize

        # fill data...
        if category_sort[c] != "Legale Pornographie":
            # timeline
            row_cells = table.add_row().cells
            row_cells[0].text = "Verteilung im Zeitraum:"
            row_cells[1].text = "{}".format(cat.getGroupedDates())
            # proportion storage <-> browser cache
            row_cells = table.add_row().cells
            row_cells[0].text = "Anteil Browsercache:"
            row_cells[1].text = "{}".format(getBrowserPercent(cat.getBrowserCacheTotal(), cat.getCounts()[0]))
        
        # format table
        for row in table.rows[1:]:
            i=-1
            row.height = table_rowheight
            for cell in row.cells:
                i+=1
                # cell alignment
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # font
                run = cell.paragraphs[0].runs[0]
                run.font.name = text_fontname
                run.font.size = table_fontsize
                if i == 0:
                    # name-column bold
                    run.font.bold = True
        document.add_paragraph().paragraph_format.space_after = Pt(0)

    counter += 1
    # update progressbar
    progress(counter, totallength)


    # write results of devices
    for d in devices:
        counter += 1
        document.add_heading(d, 2)
        for c in sorted(category_sort.keys()):
            if category_sort[c] not in devices[d].categories:
                continue

            cat = devices[d].getCategory(category_sort[c])
            # write table...
            table = document.add_table(rows=1, cols=2, style="Table Grid")
            # format header
            hdr_cells = table.rows[0].cells
            # cell merging
            hdr_cells[0].merge(hdr_cells[1])
            hdr_cells[0].text = cat.name
            # background color
            cellprop = hdr_cells[0]._tc.get_or_add_tcPr()
            cellshade = OxmlElement("w:shd")
            cellshade.set(qn("w:fill"), "#CCCCCC")
            cellprop.append(cellshade)
            # row alignment
            table.rows[0].height = table_rowheight
            hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # font
            run = hdr_cells[0].paragraphs[0].runs[0]
            run.font.name = text_fontname
            run.font.size = table_fontsize
            run.font.bold = True

            # fill data...
            # count & mediatype
            row_cells = table.add_row().cells
            row_cells[0].text = "Menge/Dateityp:"
            row_cells[1].text = cat.getCountsString()
            if category_sort[c] != "Legale Pornographie":
                # daterange
                row_cells = table.add_row().cells
                row_cells[0].text = "Erstellung auf Datenträger:"
                row_cells[1].text = "{}".format(cat.getDateRange())
                # timeline
                row_cells = table.add_row().cells
                row_cells[0].text = "Verteilung im Zeitraum:"
                row_cells[1].text = "{}".format(cat.getGroupedDates())
                # proportion storage <-> browser cache
                row_cells = table.add_row().cells
                row_cells[0].text = "Anteil Browsercache:"
                row_cells[1].text = "{}".format(getBrowserPercent(cat.getBrowserCacheTotal(), cat.getCounts()[0]))
                # paths
                row_cells = table.add_row().cells
                row_cells[0].text = "Häufigste Speicherorte:"
                # show top-paths
                rows = ""
                i = 0
                # copy the pathlist and add a thumbcache- and browsercache-entries with the total sums to the temporary copy
                temppaths = dict(cat.paths)
                thumbsum = cat.getThumbcacheSum()
                if thumbsum > 0:
                    temppaths[name_for_thumbcache] = thumbsum
                browser_sums = cat.getBrowserCacheSums()
                for b in browser_sums.keys():
                    temppaths[name_for_browsercache+" "+b] = browser_sums[b]

                # work with the temporary pathlist incl. the thumbcache-entry
                for k in sorted(temppaths, key=temppaths.get, reverse=True):
                    i += 1
                    if i > config["result"]["number_of_showed_paths"]:
                        break
                    if i > 1:
                        rows += "\n"
                    rows += "- {}".format(shortenPath(k))
                row_cells[1].text = rows
            
            # format table
            for row in table.rows[1:]:
                i=-1
                row.height = table_rowheight
                for cell in row.cells:
                    i+=1
                    # cell alignment
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # font
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = text_fontname
                    run.font.size = table_fontsize
                    if i == 0:
                        # name-column bold
                        run.font.bold = True
            document.add_paragraph().paragraph_format.space_after = Pt(0)

        progress(counter, totallength)
        document.save(result_filename)
    
def writePathDetails():
    """
    creates the outputfile (txt) with detailed information
    """
    name = config["result"]["pathdetails_name"]
    if not name.endswith(".txt"):
        name = name+".txt"
    name = f"{result_basename}_{name}"
    enc = config["result"]["pathdetails_encoding"]
    file_result = open(result_path+os.path.sep+name,"w", encoding=enc)
    # write results of file-analyze
    file_result.write("GRIFFEYE-CRAWLER - Pfad-Details vom {}\n".format(datetime.now().strftime("%d.%m.%Y")))
    file_result.write("="*47+"\n")
    file_result.write("Analysierte Datei:     {}\n".format(input_filename))
    file_result.write("Anzahl Datensätze:     {}\n".format(linecount))
    file_result.write("\n")

    # write results of devices
    counter = 0
    for d in devices:
        counter += 1
        file_result.write("\n{}\n".format(getTitleString(d, "=")))
        file_result.write("{} Dateien (Legal: {}, Illegal: {})".format(devices[d].getCounts()[0], devices[d].getCounts()[1], devices[d].getCounts()[2]))
        file_result.write("  >>  {:.2f}% illegal\n".format((devices[d].getCounts()[2]/devices[d].getCounts()[0])*100))
        for c in sorted(category_sort.keys()):
            if category_sort[c] not in devices[d].categories:
                continue

            cat = devices[d].getCategory(category_sort[c])
            file_result.write("\n{}\n".format(getTitleString(cat.name, "\u0387")))
            # count & mediatype
            file_result.write("Menge/Dateityp:\t")
            file_result.write("{}\n".format(cat.getCountsString()))

            # daterange
            file_result.write("Erstellung auf Datenträger:\t{}".format(cat.getDateRange()))
            file_result.write("\n")
            # timeline
            file_result.write("Verteilung im Zeitraum:\t{}".format(cat.getGroupedDates()))
            file_result.write("\n")
            # proportion storage <-> browser cache
            browser_total = cat.getBrowserCacheTotal()
            counts_total = cat.getCounts()[0]
            perc = (browser_total/counts_total)*100
            perc_str = "{:.0f}%".format(perc)
            if round(perc, 0) == 0 and perc > 0:
                perc_str = "<1%"
            file_result.write("Anteil Browsercache:\t{} >>> (Total: {}, Browsercache: {})".format(perc_str, counts_total, browser_total))
            file_result.write("\n")
            # paths
            file_result.write("Speicherorte:\t")
            file_result.write("\n")

            # show paths
            # copy the pathlist and add a thumbcache-entry with the total sum to the temporary copy
            temppaths = dict(cat.paths)
            thumbsum = cat.getThumbcacheSum()
            if thumbsum > 0:
                temppaths[name_for_thumbcache] = thumbsum
            # work with the temporary pathlist incl. the thumbcache-entry
            for k in sorted(temppaths, key=temppaths.get, reverse=True):
                file_result.write("- {} >>> {}\n".format(k, str(temppaths[k])))
            # if available, write other caches
            if len(cat.cachepaths)>0:
                file_result.write("    > Caches <\n")
                for k in sorted(cat.cachegroups, key=cat.cachegroups.get, reverse=True):
                    file_result.write("- {} >>> {}\n".format(k, str(cat.cachegroups[k])))
                file_result.write("    > Cache-Details <\n")
                for k in sorted(cat.cachepaths, key=cat.cachepaths.get, reverse=True):
                    file_result.write("- {} >>> {}\n".format(k, str(cat.cachepaths[k])))

        file_result.write("\n")
        # update progressbar
        progress(counter, len(devices))

    file_result.close()


# init
column_index = {}
devices = {}
column_count = 0
linecount = 0
empty_date = datetime.strptime("01.01.0001", "%d.%m.%Y")
empty_date_string = "01.01.0001 00:00:00"
unix_date = datetime.strptime("01.01.1970", "%d.%m.%Y")
invalid_lines = []
csv_separator = ""

try:
    print("===== GRIFFEYE-CRAWLER {} =====".format(version))

    # read configurations
    with open('config.json') as c:
        data = c.read()
    config = json.loads(data)
    input_encoding = config["input"]["encoding"]
    result_filename = ""
    result_encoding = config["result"]["encoding"]
    category_legality = {}
    category_sort = {}
    for cat in config["categories"]:
        category_legality[cat["name"]] = cat["legality"]
        category_sort[cat["sort"]] = cat["name"]
    known_cache_paths = {}
    browser_names = []
    thumbcache_names = []
    for cac in config["caches"]:
        known_cache_paths[cac["path"]] = cac["name"]
        if cac["is_browser"] and cac["name"] not in browser_names:
            browser_names.append(cac["name"])
        if cac["is_thumbcache"] and cac["name"] not in thumbcache_names:
            thumbcache_names.append(cac["name"])

    # ask for informations
    input_filename = input("Pfad/Name des Input-CSV > ")
    # remove " & ' from path (prevents error while reading the file)
    input_filename = input_filename.replace("\"", "")
    input_filename = input_filename.replace("'", "")
    
    default_format = "docx"
    result_format = input("Format des Ergebnisses (Default: {}) [txt, docx] > ".format(default_format)) or default_format
    if result_format.strip().lower() == "docx":
        result_format = "docx"
    if result_format.strip().lower() == "txt":
        result_format = "txt"
    print()

    # generate result path/name based on inputfile
    result_path = os.path.dirname(input_filename)
    temp_filename = os.path.basename(input_filename)
    result_basename = os.path.splitext(temp_filename)[0]
    result_filename = result_path+os.path.sep+result_basename+"."+result_format

    # get linecount for progressbar
    linecount = getLinecount(input_filename)

    # analyze file
    print("Analysiere Datei '{}'...".format(input_filename))
    analyzeFile(input_filename)
    if len(invalid_lines) > 0:
        print()
        print("  !!! Ungültige Zeilen in Input-CSV entdeckt und in Verarbeitung ignoriert")
        print("  !!! Zeilen: ", end="")
        for l in invalid_lines:
            print(l, end="  ")
        print()
    print()

    # process data
    print("Verarbeite Datensätze...")
    file_input = open(input_filename, "r", encoding=input_encoding)
    result = ""
    counter = -1
    for line in file_input:
        counter += 1
        # ignore csv-header
        if counter == 0:
            continue

        # get data from file
        try:
            if line.count(csv_separator) != column_count:
                column = convertLine(line, counter+1)
            else:
                column = line.split(csv_separator)
            data_category = column[column_index['col_category']]
            data_path = column[column_index['col_path']]
            data_type = column[column_index['col_type']]
            data_date = column[column_index['col_date']]
            # if date is empty (01.01.0001) try the alternative date
            if data_date == empty_date_string and column_index['col_alt_date'] > -1:
                data_date = column[column_index['col_alt_date']]
            data_device = column[column_index['col_device']]
            data_hash = column[column_index['col_hash']]
            # add data to device
            device = devices[data_device]
            device.addFile(data_category, data_path, data_type, data_date, data_hash)
        except LineNotValidException as exp:
            pass
        # update progressbar
        progress(counter, linecount)
    file_input.close()

    # generate total from devices
    total = {}
    cat_devcount = {}
    for d in devices:
        categories = devices[d].getCategories()
        for dev_cat in categories.values():
            # increase/generate devicecount for category
            if dev_cat.name not in cat_devcount.keys():
                cat_devcount[dev_cat.name] = 1
            else:
                cat_devcount[dev_cat.name] += 1
            # get/generate total category
            total_cat = None
            if dev_cat.name not in total.keys():
                total_cat = Category(dev_cat.name, dev_cat.min_date)
                total[dev_cat.name] = total_cat
            else:
                total_cat = total[dev_cat.name]
            # merge device-category to total-category
            total_cat.merge(dev_cat)
    print()

    # write output-files
    print("Schreibe Ergebnisdatei...")
    name_for_thumbcache = config["other"]["name_for_thumbcache"]
    name_for_browsercache = config["other"]["name_for_browsercache"]
    
    if result_format == "txt":
        writeOutputfileTxt()
    elif result_format == "docx":
        writeOutputfileDocx()
    else:
        # actually not possible...
        raise ResultFormatUnknownException(result_format)
    if config["result"]["generate_pathdetails"]:
        writePathDetails()

    print()
    print()
    print("ERLEDIGT! {} Datensätze verarbeitet (siehe '{}')".format(counter, result_filename))

except ColumnNotFoundException as exp:
    print()
    print("ERROR: Verarbeitung abgebrochen!")
    print(">", exp.message)
except ResultFormatUnknownException as exp:
    print()
    print("ERROR: Verarbeitung abgebrochen!")
    print(">", exp.message)
except FileNotFoundError as exp:
    print()
    print("ERROR: Verarbeitung abgebrochen!")
    print("> Datei '{}' nicht gefunden".format(exp.filename))
except KeyError as exp:
    print()
    print("ERROR: Verarbeitung abgebrochen!")
    print("> Konfiguration '{}' nicht gefunden".format(exp))
except UnicodeDecodeError as exp:
    print()
    print("ERROR: Verarbeitung abgebrochen!")
    if exp.args[0] == "utf-8":
        print("Datei liegt nicht im UTF-8-Format vor. Config anpassen oder Datei umwandeln...")
    else:
        print("Datei liegt in einem unbekannten Format vor")
except UnicodeError as exp:
    print()
    print("ERROR: Verarbeitung abgebrochen!")
    if "UTF-16" in exp.args[0]:
        print("Datei liegt nicht im UTF-16-Format vor. Config anpassen oder Datei umwandeln...")
    else:
        print("Datei liegt in einem unbekannten Format vor")
except Exception as exp:
    print()
    print("ERROR: Verarbeitung abgebrochen!")
    traceback.print_exc()

print()
input()
